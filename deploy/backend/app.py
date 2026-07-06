from __future__ import annotations

import csv
import io
import os
import re
import uuid
from datetime import datetime, timedelta
from json import dumps
from typing import Any

from fastapi import Depends, FastAPI, File, Header, HTTPException, Request, UploadFile
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse, JSONResponse, Response, StreamingResponse

# 上传文件落盘目录(可用 UPLOAD_DIR 覆盖);按租户分目录。
UPLOAD_DIR = os.path.abspath(os.getenv("UPLOAD_DIR", os.path.join(os.path.dirname(__file__), "..", "uploads")))
MAX_UPLOAD_BYTES = 50 * 1024 * 1024  # 50MB
from pydantic import BaseModel, Field, field_validator

# 数据校验共用枚举/常量(与 DB enum 对齐)。
_FUND_STATUS = {"planning", "raising", "investing", "harvesting", "extended", "closed"}

import auth
import llm
from auth import AuthedUser, create_token, current_user, require_permission, require_roles, verify_password

try:
    import pymysql
except Exception:  # pragma: no cover - dependency is installed by deploy.sh
    pymysql = None


APP_NAME = "CapitalOS Backend"
APP_VERSION = "1.0.0"


class LoginPayload(BaseModel):
    account: str
    password: str


class GenericActionPayload(BaseModel):
    action: str
    entity_type: str
    entity_id: int | None = None
    entity_label: str | None = None
    before: dict[str, Any] | None = None
    after: dict[str, Any] | None = None
    risk_level: str = "low"


class CreateProjectPayload(BaseModel):
    short_name: str = Field(min_length=1, max_length=120)
    legal_name: str | None = Field(default=None, max_length=200)
    industry_group: str = Field(default="Enterprise Software", max_length=120)
    city: str = Field(default="Shanghai", max_length=80)
    summary: str | None = Field(default=None, max_length=2000)

    @field_validator("short_name")
    @classmethod
    def _nonblank_name(cls, v: str) -> str:
        v = v.strip()
        if not v:
            raise ValueError("项目名称不能为空")
        return v


class CreateFundPayload(BaseModel):
    fund_name: str = Field(min_length=1, max_length=180)
    legal_name: str | None = Field(default=None, max_length=200)
    target_size: float = Field(default=100000000, ge=0)
    committed_size: float = Field(default=0, ge=0)
    paid_in_size: float = Field(default=0, ge=0)
    fund_status: str = "planning"

    @field_validator("fund_name")
    @classmethod
    def _nonblank_name(cls, v: str) -> str:
        v = v.strip()
        if not v:
            raise ValueError("基金名称不能为空")
        return v

    @field_validator("fund_status")
    @classmethod
    def _valid_status(cls, v: str) -> str:
        if v not in _FUND_STATUS:
            raise ValueError(f"fund_status 必须是 {sorted(_FUND_STATUS)} 之一")
        return v


class CreateDocumentPayload(BaseModel):
    title: str
    document_kind: str = "shared"
    file_name: str = "frontend-upload-placeholder.pdf"
    storage_uri: str = "mock://documents/frontend-upload-placeholder.pdf"
    file_size_bytes: int = 1024


class StartWorkflowPayload(BaseModel):
    title: str
    workflow_family: str = "project"
    payload: dict[str, Any] | None = None


class WorkflowActionPayload(BaseModel):
    action: str
    comment: str | None = None


class RiskUpdatePayload(BaseModel):
    update_text: str
    update_status: str = "progress"


class ImportExportPayload(BaseModel):
    task_kind: str
    entity_type: str
    source_file_uri: str | None = None


class PreferencePayload(BaseModel):
    notification_json: dict[str, Any] | None = None
    favorite_nav_json: list[str] | None = None
    table_view_json: dict[str, Any] | None = None


class UpdateProjectPayload(BaseModel):
    short_name: str | None = Field(default=None, min_length=1, max_length=120)
    legal_name: str | None = Field(default=None, max_length=200)
    industry_group: str | None = Field(default=None, max_length=120)
    city: str | None = Field(default=None, max_length=80)
    summary: str | None = Field(default=None, max_length=2000)
    stage_label: str | None = Field(default=None, max_length=80)
    # 基本情况分组扩展字段(项目卡片二级 tab「基本情况」可编辑)
    registered_location: str | None = Field(default=None, max_length=200)
    registry_code_mask: str | None = Field(default=None, max_length=120)
    source_channel: str | None = Field(default=None, max_length=120)
    thesis: str | None = Field(default=None, max_length=2000)
    product_note: str | None = Field(default=None, max_length=2000)
    highlight_note: str | None = Field(default=None, max_length=2000)
    # 乐观锁:前端把加载时的 updated_at 回传;与库中不一致 → 409(有人先改了)。
    expected_updated_at: str | None = None


class UpdateFundPayload(BaseModel):
    fund_name: str | None = Field(default=None, min_length=1, max_length=180)
    legal_name: str | None = Field(default=None, max_length=200)
    fund_status: str | None = None
    target_size: float | None = Field(default=None, ge=0)
    committed_size: float | None = Field(default=None, ge=0)
    paid_in_size: float | None = Field(default=None, ge=0)
    net_asset_value: float | None = Field(default=None, ge=0)
    expected_updated_at: str | None = None

    @field_validator("fund_status")
    @classmethod
    def _valid_status(cls, v: str | None) -> str | None:
        if v is not None and v not in _FUND_STATUS:
            raise ValueError(f"fund_status 必须是 {sorted(_FUND_STATUS)} 之一")
        return v


class ImportCsvPayload(BaseModel):
    csv_text: str


class BpParsePayload(BaseModel):
    text: str


class AiSummarizePayload(BaseModel):
    text: str
    kind: str = "meeting"  # meeting | research


class AiAnalyzePayload(BaseModel):
    text: str
    instruction: str | None = None


class AiConfigPayload(BaseModel):
    enabled: bool | None = None
    base_url: str | None = None
    api_key: str | None = None  # 空/缺省 = 保留原 key 不变
    model: str | None = None
    timeout: int | None = None


class DelegationPayload(BaseModel):
    delegatee_user_id: int
    workflow_family: str = "all"  # project | fund | office | all
    starts_at: str | None = None  # ISO / YYYY-MM-DD;缺省=现在
    ends_at: str | None = None    # 缺省=+30 天
    reason: str | None = None


SCREENS: list[dict[str, str]] = [
    {"id": "login", "title": "Login", "group": "Entry"},
    {"id": "workbench", "title": "Executive Workbench", "group": "Workspace"},
    {"id": "ai-workspace", "title": "AI Workspace", "group": "Workspace"},
    {"id": "announcements", "title": "Announcements", "group": "Collaboration"},
    {"id": "calendar", "title": "Calendar", "group": "Collaboration"},
    {"id": "message-center", "title": "Message Center", "group": "Collaboration"},
    {"id": "flow-center", "title": "Workflow Center", "group": "Workflow"},
    {"id": "flow-project", "title": "Project Workflows", "group": "Workflow"},
    {"id": "flow-fund", "title": "Fund Workflows", "group": "Workflow"},
    {"id": "flow-oa", "title": "Office Workflows", "group": "Workflow"},
    {"id": "project-board", "title": "Project Board", "group": "Projects"},
    {"id": "project-list", "title": "Project List", "group": "Projects"},
    {"id": "project-add", "title": "Add Project", "group": "Projects"},
    {"id": "project-detail-overview", "title": "Project Detail Overview", "group": "Projects"},
    {"id": "project-detail-investment", "title": "Project Investment Relation", "group": "Projects"},
    {"id": "project-detail-postdata", "title": "Project Post Data", "group": "Projects"},
    {"id": "meeting-ai", "title": "Meeting AI", "group": "Projects"},
    {"id": "fund-list", "title": "Fund List", "group": "Funds"},
    {"id": "fund-add", "title": "Add Fund", "group": "Funds"},
    {"id": "fund-detail-overview", "title": "Fund Detail Overview", "group": "Funds"},
    {"id": "fund-detail-cashflow", "title": "Fund Cashflow", "group": "Funds"},
    {"id": "fund-detail-financials", "title": "Fund Financials", "group": "Funds"},
    {"id": "investment-info", "title": "Investment Ledger", "group": "Funds"},
    {"id": "equity-change", "title": "Equity Change", "group": "Funds"},
    {"id": "investor-list", "title": "Investor List", "group": "Investors"},
    {"id": "investor-detail", "title": "Investor Detail", "group": "Investors"},
    {"id": "manager-orgs", "title": "Management Organizations", "group": "Investors"},
    {"id": "post-data-collection", "title": "Post Data Collection", "group": "Portfolio"},
    {"id": "risk-clauses", "title": "Key Clauses", "group": "Risk"},
    {"id": "burst-risk", "title": "Risk Incidents", "group": "Risk"},
    {"id": "document-center", "title": "Document Center", "group": "Documents"},
    {"id": "process-files", "title": "Workflow Files", "group": "Documents"},
    {"id": "research-library", "title": "Research Library", "group": "AI Data"},
    {"id": "internal-research", "title": "Internal Research", "group": "AI Data"},
    {"id": "report-dashboard", "title": "Report Dashboard", "group": "Reports"},
    {"id": "import-export", "title": "Import Export Center", "group": "Common"},
    {"id": "system-users", "title": "Users And Organizations", "group": "System"},
    {"id": "roles-permissions", "title": "Roles And Permissions", "group": "System"},
    {"id": "field-config", "title": "Fields And Forms", "group": "System"},
    {"id": "account-settings", "title": "Account Settings", "group": "Account"},
    {"id": "recycle-bin", "title": "Recycle Bin", "group": "System"},
]

PROJECTS: list[dict[str, Any]] = [
    {
        "id": 1,
        "name": "Matrix Medical",
        "stage": "Project Approval",
        "sector": "Medical Devices",
        "city": "Shanghai",
        "owner": "Nina Lin",
        "risk": "medium",
        "next_step": "Prepare IC material",
    },
    {
        "id": 2,
        "name": "Northstar Storage",
        "stage": "TS",
        "sector": "Energy Storage",
        "city": "Changzhou",
        "owner": "Nina Lin",
        "risk": "medium",
        "next_step": "Update TS clauses",
    },
    {
        "id": 3,
        "name": "Lanzhou Robotics",
        "stage": "Diligence",
        "sector": "Robotics",
        "city": "Shenzhen",
        "owner": "Nina Lin",
        "risk": "low",
        "next_step": "Finish legal checklist",
    },
]

FUNDS: list[dict[str, Any]] = [
    {"id": 1, "name": "Growth Fund I", "status": "investing", "committed_size": 1860000000, "tvpi": 1.31},
    {"id": 2, "name": "Carbon Fund I", "status": "investing", "committed_size": 2000000000, "tvpi": 1.12},
    {"id": 3, "name": "Healthcare Special Fund", "status": "investing", "committed_size": 1200000000, "tvpi": 1.10},
]

RISKS: list[dict[str, Any]] = [
    {"id": 1, "title": "Cash balance below safety line", "project": "Starfield Agri", "severity": "high"},
    {"id": 2, "title": "Related-party disclosure incomplete", "project": "Qingqiong Chip", "severity": "high"},
    {"id": 3, "title": "Clinical milestone delay", "project": "Matrix Medical", "severity": "medium"},
]

# 租户作用域的 ledger:每条查询内含一个 tenant_id=%s 占位,由 screen_ledger 注入。
# 未列入 LEDGER_TENANT_SCOPED 的(角色/字段/用户等全局配置)不带该参数。
LEDGER_QUERIES: dict[str, str] = {
    "announcements": """
        SELECT title AS '标题', publish_status AS '状态', published_at AS '发布时间',
               expires_at AS '到期时间', created_at AS '创建时间'
        FROM cap_announcements
        WHERE deleted_at IS NULL AND tenant_id=%s
        ORDER BY announcement_id DESC
        """,
    "calendar": """
        SELECT event_title AS '日程', event_kind AS '类型', linked_entity_type AS '关联类型',
               starts_at AS '开始时间', ends_at AS '结束时间', location_text AS '地点',
               visibility AS '可见范围'
        FROM cap_calendar_events
        WHERE deleted_at IS NULL AND tenant_id=%s
        ORDER BY starts_at DESC
        """,
    "message-center": """
        SELECT source_kind AS '类型', title AS '标题', message_box AS '收件箱',
               action_status AS '处理状态', created_at AS '创建时间'
        FROM cap_messages
        WHERE tenant_id=%s
        ORDER BY message_id DESC
        """,
    "project-list": """
        SELECT p.project_id AS '__id', p.short_name AS '项目名称', p.stage_label AS '阶段', p.opportunity_status AS '状态',
               COALESCE(u.display_name, '-') AS '负责人', p.city AS '城市',
               p.industry_group AS '行业方向', p.updated_at AS '更新时间'
        FROM cap_projects p
        LEFT JOIN cap_users u ON u.user_id=p.owner_user_id
        WHERE p.deleted_at IS NULL AND p.tenant_id=%s
        ORDER BY p.project_id DESC
        """,
    "fund-list": """
        SELECT f.fund_id AS '__id', f.fund_name AS '基金简称', f.fund_status AS '状态', f.raise_method AS '募集方式',
               f.committed_size AS '认缴规模', f.paid_in_size AS '实缴总额',
               f.net_asset_value AS '净资产', f.unit_nav AS '单位净值',
               COALESCE(m.org_name, '-') AS '管理人'
        FROM cap_funds f
        LEFT JOIN cap_management_orgs m ON m.management_org_id=f.manager_org_id
        WHERE f.deleted_at IS NULL AND f.tenant_id=%s
        ORDER BY f.fund_id DESC
        """,
    "investment-info": """
        SELECT COALESCE(p.short_name, '-') AS '项目名称', COALESCE(f.fund_name, '-') AS '投资主体',
               i.round_label AS '轮次', i.agreement_amount AS '协议金额',
               i.paid_amount AS '累计打款', i.shareholding_pct AS '最新持股',
               i.position_status AS '状态'
        FROM cap_investment_positions i
        LEFT JOIN cap_projects p ON p.project_id=i.project_id
        LEFT JOIN cap_funds f ON f.fund_id=i.fund_id
        WHERE i.tenant_id=%s
        ORDER BY i.investment_position_id DESC
        """,
    "equity-change": """
        SELECT COALESCE(p.short_name, '-') AS '项目名称', e.change_reason AS '变更原因',
               e.round_label AS '轮次', e.before_pct AS '交易前股比',
               e.after_pct AS '交易后股比', e.signed_on AS '协议时间',
               e.created_at AS '创建时间'
        FROM cap_equity_changes e
        LEFT JOIN cap_projects p ON p.project_id=e.project_id
        WHERE e.tenant_id=%s
        ORDER BY e.equity_change_id DESC
        """,
    "investor-list": """
        SELECT investor_name AS '投资人', investor_kind AS '类型', qualification_status AS '适当性',
               risk_rating AS '风险评级', city AS '城市', disclosure_status AS '披露状态',
               updated_at AS '更新时间'
        FROM cap_investors
        WHERE deleted_at IS NULL AND tenant_id=%s
        ORDER BY investor_id DESC
        """,
    "manager-orgs": """
        SELECT org_name AS '机构名称', org_kind AS '类型', city AS '城市',
               contact_name AS '联系人', status AS '状态', updated_at AS '更新时间'
        FROM cap_management_orgs
        WHERE deleted_at IS NULL AND tenant_id=%s
        ORDER BY management_org_id DESC
        """,
    "post-data-collection": """
        SELECT campaign_name AS '收集任务', period_code AS '期间', status AS '状态',
               due_on AS '截止日期', frequency AS '频率', send_mode AS '发送方式',
               created_at AS '创建时间'
        FROM cap_data_collection_campaigns
        WHERE tenant_id=%s
        ORDER BY collection_campaign_id DESC
        """,
    "research-library": """
        SELECT note_title AS '标题', note_kind AS '类型', source_name AS '来源',
               review_status AS '状态', updated_at AS '更新时间'
        FROM cap_research_notes
        WHERE deleted_at IS NULL AND tenant_id=%s
        ORDER BY research_note_id DESC
        """,
    "internal-research": """
        SELECT note_title AS '标题', note_kind AS '类型', source_name AS '来源',
               review_status AS '状态', updated_at AS '更新时间'
        FROM cap_research_notes
        WHERE deleted_at IS NULL AND tenant_id=%s
        ORDER BY research_note_id DESC
        """,
    "import-export": """
        SELECT task_code AS '任务编号', task_kind AS '类型', entity_type AS '模块',
               task_status AS '状态', total_rows AS '总行数', success_rows AS '成功',
               failed_rows AS '失败', requested_at AS '创建时间'
        FROM cap_import_export_tasks
        WHERE tenant_id=%s
        ORDER BY import_export_task_id DESC
        """,
    "system-users": """
        SELECT u.display_name AS '姓名', COALESCE(o.org_name, '-') AS '部门',
               u.login_name AS '账号', u.email AS '邮箱', u.account_status AS '状态',
               u.last_login_at AS '最近登录'
        FROM cap_users u
        LEFT JOIN cap_organizations o ON o.org_id=u.org_id
        WHERE u.deleted_at IS NULL
        ORDER BY u.user_id DESC
        """,
    "roles-permissions": """
        SELECT role_name AS '角色', data_scope AS '数据范围', is_system_role AS '系统角色',
               description AS '描述', updated_at AS '更新时间'
        FROM cap_roles
        ORDER BY role_id DESC
        """,
    "field-config": """
        SELECT entity_type AS '所属对象', field_label AS '字段名称', data_type AS '类型',
               is_required AS '必填', is_searchable AS '可搜索', display_order AS '排序',
               updated_at AS '更新时间'
        FROM cap_custom_field_definitions
        WHERE deleted_at IS NULL
        ORDER BY entity_type, display_order
        """,
    "recycle-bin": """
        SELECT recycle_item_id AS '__id', object_label AS '对象', object_type AS '类型', delete_reason AS '删除原因',
               purge_status AS '状态', deleted_at AS '删除时间',
               recoverable_until AS '保留到期'
        FROM cap_recycle_items
        WHERE tenant_id=%s AND purge_status <> 'purged'
        ORDER BY recycle_item_id DESC
        """,
}

# 需注入 tenant_id=%s 的 ledger 屏(其余为全局配置屏:角色/字段/用户列表)。
LEDGER_TENANT_SCOPED: set[str] = {
    "announcements", "calendar", "message-center", "project-list", "fund-list",
    "investment-info", "equity-change", "investor-list", "manager-orgs",
    "post-data-collection", "research-library", "internal-research",
    "import-export", "recycle-bin",
}


app = FastAPI(title=APP_NAME, version=APP_VERSION)

# CORS:鉴权走 Bearer token(非 cookie),故 allow_credentials=False —— 这样 `*` 也合规,
# 且可通过 CORS_ALLOW_ORIGINS 收紧到固定域名(生产建议填具体源)。
app.add_middleware(
    CORSMiddleware,
    allow_origins=os.getenv("CORS_ALLOW_ORIGINS", "*").split(","),
    allow_credentials=False,
    allow_methods=["*"],
    allow_headers=["*"],
)

# 无需鉴权的路径:健康检查、登录本身。其余 /api/* 一律要求有效 Bearer token。
_PUBLIC_PATHS = {"/health", "/api/auth/login", "/api/db/ping"}

# 简易滑动窗口限流(进程内):登录防暴破、AI 端点防刷/控成本。
import time as _time
from collections import defaultdict as _defaultdict, deque as _deque

_RATE_BUCKETS: dict[str, "_deque[float]"] = _defaultdict(_deque)
# (路径前缀, 窗口秒, 上限次数)
_RATE_RULES = [
    ("/api/auth/login", 60.0, 10),   # 每 IP 每分钟最多 10 次登录尝试
    ("/api/ai/", 60.0, 30),          # 每 IP 每分钟最多 30 次 AI 调用
]


def _client_ip(request: Request) -> str:
    fwd = request.headers.get("x-forwarded-for", "")
    return fwd.split(",")[0].strip() if fwd else (request.client.host if request.client else "unknown")


def _rate_allow(key: str, window: float, limit: int) -> bool:
    now = _time.monotonic()
    dq = _RATE_BUCKETS[key]
    while dq and dq[0] <= now - window:
        dq.popleft()
    if len(dq) >= limit:
        return False
    dq.append(now)
    return True


@app.middleware("http")
async def rate_limit(request: Request, call_next):
    path = request.url.path
    if request.method != "OPTIONS":
        for prefix, window, limit in _RATE_RULES:
            if path.startswith(prefix):
                if not _rate_allow(f"{prefix}|{_client_ip(request)}", window, limit):
                    origin = request.headers.get("origin", "*")
                    return JSONResponse(
                        {"detail": "请求过于频繁,请稍后再试"},
                        status_code=429,
                        headers={"Access-Control-Allow-Origin": origin, "Retry-After": "60", "Vary": "Origin"},
                    )
                break
    return await call_next(request)


@app.middleware("http")
async def enforce_auth(request: Request, call_next):
    """兜底鉴权闸门:任何 /api 业务路径无有效 token 直接 401 —— 保证没有端点被漏掉。
    身份归属仍由各写端点的 Depends(current_user) 从同一 token 派生。"""
    path = request.url.path
    if request.method == "OPTIONS" or path in _PUBLIC_PATHS or not path.startswith("/api"):
        return await call_next(request)
    authz = request.headers.get("authorization", "")
    origin = request.headers.get("origin", "*")

    def _fail(status: int, detail: str) -> JSONResponse:
        # 跨源 401 也要带 CORS 头,否则浏览器读不到状态码。
        return JSONResponse(
            {"detail": detail},
            status_code=status,
            headers={"Access-Control-Allow-Origin": origin, "Vary": "Origin"},
        )

    if not authz.lower().startswith("bearer "):
        return _fail(401, "Missing bearer token")
    try:
        auth.decode_token(authz.split(" ", 1)[1].strip())
    except HTTPException as exc:
        return _fail(exc.status_code, str(exc.detail))
    return await call_next(request)


def db_config() -> dict[str, Any]:
    host = os.getenv("DB_HOST")
    user = os.getenv("DB_USER")
    password = os.getenv("DB_PASSWORD")
    database = os.getenv("DB_NAME", "capitalos")
    port = int(os.getenv("DB_PORT", "3306"))

    if not host or not user or password is None:
        raise HTTPException(status_code=400, detail="DB_HOST, DB_USER and DB_PASSWORD are required")

    return {
        "host": host,
        "port": port,
        "user": user,
        "password": password,
        "database": database,
        "charset": "utf8mb4",
        "cursorclass": pymysql.cursors.DictCursor,
        "connect_timeout": 5,
        "autocommit": False,
    }


def connect_db():
    if pymysql is None:
        raise HTTPException(status_code=500, detail="pymysql is not installed")
    return pymysql.connect(**db_config())


def resolve_tenant_id(cursor: Any, org_id: int | None) -> int | None:
    """从用户组织向上走 parent_org_id 到根公司(parent 为 NULL),返回其 org_id 作为租户。
    深度设上限防脏数据成环。"""
    current = org_id
    for _ in range(16):
        if current is None:
            return None
        cursor.execute("SELECT parent_org_id FROM cap_organizations WHERE org_id=%s", (current,))
        row = cursor.fetchone()
        if row is None:
            return current
        parent = row["parent_org_id"]
        if parent is None:
            return current  # 已是根公司
        current = parent
    return current


def tenant_of(user: "AuthedUser") -> int:
    """取当前请求的租户 id;缺失即拒绝(绝不放行无租户上下文的数据访问)。"""
    if user.tenant_id is None:
        raise HTTPException(status_code=403, detail="No tenant context on this account")
    return int(user.tenant_id)


def _tenant_for_user(cursor: Any, actor_user_id: int | None) -> int | None:
    """写操作盖章用:从 actor(已由 token 认证,可信)反查其租户。
    用户只在自己租户内活动,故 actor 的租户即为落库租户。"""
    if actor_user_id is None:
        return None
    cursor.execute("SELECT org_id FROM cap_users WHERE user_id=%s", (actor_user_id,))
    row = cursor.fetchone()
    return resolve_tenant_id(cursor, row["org_id"]) if row else None


def write_audit(
    cursor: Any,
    actor_user_id: int | None,
    action_code: str,
    entity_type: str,
    entity_id: int | None,
    entity_label: str | None,
    before: dict[str, Any] | None = None,
    after: dict[str, Any] | None = None,
    risk_level: str = "low",
) -> int:
    cursor.execute(
        """
        INSERT INTO cap_audit_logs
          (actor_user_id, tenant_id, action_code, entity_type, entity_id, entity_label, request_id, ip_mask,
           before_json, after_json, risk_level, occurred_at)
        VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s)
        """,
        (
            actor_user_id,
            _tenant_for_user(cursor, actor_user_id),
            action_code,
            entity_type,
            entity_id,
            entity_label,
            f"REQ-{uuid.uuid4().hex[:12]}",
            "api-client",
            dumps(before or {}, ensure_ascii=False),
            dumps(after or {}, ensure_ascii=False),
            risk_level,
            datetime.now(),
        ),
    )
    return int(cursor.lastrowid)


def ensure_ui_action_table(cursor: Any) -> None:
    cursor.execute(
        """
        CREATE TABLE IF NOT EXISTS cap_ui_action_events (
          ui_action_event_id BIGINT UNSIGNED NOT NULL AUTO_INCREMENT,
          audit_log_id BIGINT UNSIGNED NULL,
          actor_user_id BIGINT UNSIGNED NULL,
          tenant_id BIGINT UNSIGNED NULL,
          action_code VARCHAR(120) NOT NULL,
          entity_type VARCHAR(80) NOT NULL,
          entity_id BIGINT UNSIGNED NULL,
          entity_label VARCHAR(220) NULL,
          payload_json JSON NULL,
          result_json JSON NULL,
          event_status ENUM('succeeded','failed') NOT NULL DEFAULT 'succeeded',
          occurred_at DATETIME NOT NULL,
          created_at TIMESTAMP NOT NULL DEFAULT CURRENT_TIMESTAMP,
          PRIMARY KEY (ui_action_event_id),
          KEY idx_cap_ui_action_audit (audit_log_id),
          KEY idx_cap_ui_action_actor_time (actor_user_id, occurred_at),
          KEY idx_cap_ui_action_code (action_code),
          CONSTRAINT fk_cap_ui_action_audit
            FOREIGN KEY (audit_log_id) REFERENCES cap_audit_logs (audit_log_id)
            ON DELETE SET NULL,
          CONSTRAINT fk_cap_ui_action_actor
            FOREIGN KEY (actor_user_id) REFERENCES cap_users (user_id)
            ON DELETE SET NULL
        ) ENGINE=InnoDB DEFAULT CHARSET=utf8mb4 COLLATE=utf8mb4_unicode_ci
        """
    )


def write_ui_action_event(
    cursor: Any,
    actor_user_id: int | None,
    audit_id: int | None,
    payload: GenericActionPayload,
    result: dict[str, Any],
) -> int:
    ensure_ui_action_table(cursor)
    cursor.execute(
        """
        INSERT INTO cap_ui_action_events
          (audit_log_id, actor_user_id, tenant_id, action_code, entity_type, entity_id, entity_label,
           payload_json, result_json, event_status, occurred_at)
        VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s, 'succeeded', %s)
        """,
        (
            audit_id,
            actor_user_id,
            _tenant_for_user(cursor, actor_user_id),
            payload.action,
            payload.entity_type,
            payload.entity_id,
            payload.entity_label,
            dumps(payload.model_dump(), ensure_ascii=False),
            dumps(result, ensure_ascii=False),
            datetime.now(),
        ),
    )
    return int(cursor.lastrowid)


def create_fund_record(cursor: Any, payload: CreateFundPayload, actor_user_id: int | None) -> dict[str, Any]:
    allowed_status = {"planning", "raising", "investing", "harvesting", "extended", "closed"}
    status = payload.fund_status if payload.fund_status in allowed_status else "planning"
    cursor.execute(
        """
        INSERT INTO cap_funds
          (fund_code, fund_name, legal_name, fund_status, raise_method, base_currency,
           target_size, committed_size, paid_in_size, net_asset_value, created_by, tenant_id)
        VALUES (%s, %s, %s, %s, 'private', 'CNY', %s, %s, %s, %s, %s, %s)
        """,
        (
            f"FUN-API-{uuid.uuid4().hex[:8].upper()}",
            payload.fund_name,
            payload.legal_name or payload.fund_name,
            status,
            payload.target_size,
            payload.committed_size,
            payload.paid_in_size,
            payload.paid_in_size,
            actor_user_id,
            _tenant_for_user(cursor, actor_user_id),
        ),
    )
    fund_id = int(cursor.lastrowid)
    audit_id = write_audit(
        cursor,
        actor_user_id,
        "fund.create",
        "fund",
        fund_id,
        payload.fund_name,
        after=payload.model_dump(),
    )
    return {"ok": True, "fund_id": fund_id, "audit_id": audit_id, "affected_table": "cap_funds"}


def create_document_record(cursor: Any, payload: CreateDocumentPayload, actor_user_id: int | None) -> dict[str, Any]:
    allowed_kinds = {"project", "fund", "investor", "workflow", "shared", "meeting", "risk", "research", "other"}
    document_kind = payload.document_kind if payload.document_kind in allowed_kinds else "shared"
    cursor.execute(
        """
        INSERT INTO cap_documents
          (document_code, title, document_kind, storage_uri, file_name, mime_type,
           file_size_bytes, current_version_no, access_level, watermark_policy,
           fulltext_status, uploaded_by, tenant_id)
        VALUES (%s, %s, %s, %s, %s, 'application/pdf', %s, 1, 'team', 'viewer', 'queued', %s, %s)
        """,
        (
            f"DOC-API-{uuid.uuid4().hex[:8].upper()}",
            payload.title,
            document_kind,
            payload.storage_uri,
            payload.file_name,
            payload.file_size_bytes,
            actor_user_id,
            _tenant_for_user(cursor, actor_user_id),
        ),
    )
    document_id = int(cursor.lastrowid)
    cursor.execute(
        """
        INSERT INTO cap_document_versions
          (document_id, version_no, storage_uri, file_size_bytes, checksum_hash, change_note, uploaded_by, uploaded_at)
        VALUES (%s, 1, %s, %s, %s, 'Initial frontend upload', %s, %s)
        """,
        (
            document_id,
            payload.storage_uri,
            payload.file_size_bytes,
            uuid.uuid4().hex,
            actor_user_id,
            datetime.now(),
        ),
    )
    version_id = int(cursor.lastrowid)
    cursor.execute(
        """
        INSERT INTO cap_document_permissions
          (document_id, grantee_kind, grantee_id, can_preview, can_download, can_edit,
           can_delete, can_share, watermark_required, granted_by)
        VALUES (%s, 'user', %s, 1, 1, 1, 0, 1, 1, %s)
        """,
        (document_id, actor_user_id or 1, actor_user_id),
    )
    audit_id = write_audit(
        cursor,
        actor_user_id,
        "document.upload",
        "document",
        document_id,
        payload.title,
        after=payload.model_dump(),
    )
    return {
        "ok": True,
        "document_id": document_id,
        "version_id": version_id,
        "audit_id": audit_id,
        "affected_table": "cap_documents",
    }


def start_workflow_record(cursor: Any, payload: StartWorkflowPayload, actor_user_id: int | None) -> dict[str, Any]:
    family = payload.workflow_family if payload.workflow_family in {"project", "fund", "office"} else "project"
    cursor.execute(
        """
        SELECT workflow_template_id, template_name
        FROM cap_workflow_templates
        WHERE workflow_family=%s
        ORDER BY workflow_template_id
        LIMIT 1
        """,
        (family,),
    )
    template = cursor.fetchone()
    if template is None:
        cursor.execute("SELECT workflow_template_id, template_name FROM cap_workflow_templates ORDER BY workflow_template_id LIMIT 1")
        template = cursor.fetchone()
    if template is None:
        raise HTTPException(status_code=400, detail="No workflow template exists")

    cursor.execute(
        """
        INSERT INTO cap_workflow_instances
          (workflow_template_id, instance_code, title, initiator_user_id, instance_status,
           current_step_key, payload_json, started_at, tenant_id)
        VALUES (%s, %s, %s, %s, 'running', 'start', %s, %s, %s)
        """,
        (
            template["workflow_template_id"],
            f"WF-API-{uuid.uuid4().hex[:8].upper()}",
            payload.title,
            actor_user_id,
            dumps(payload.payload or {}, ensure_ascii=False),
            datetime.now(),
            _tenant_for_user(cursor, actor_user_id),
        ),
    )
    instance_id = int(cursor.lastrowid)
    # 取模板第一步(sort_order 最小),建首步任务;若模板没配步骤则回落到笼统首节点。
    cursor.execute(
        """
        SELECT workflow_step_id, step_key, step_name
        FROM cap_workflow_steps WHERE workflow_template_id=%s
        ORDER BY sort_order, workflow_step_id LIMIT 1
        """,
        (template["workflow_template_id"],),
    )
    first_step = cursor.fetchone()
    step_id = first_step["workflow_step_id"] if first_step else None
    step_name = first_step["step_name"] if first_step else "首节点审批"
    if first_step:
        cursor.execute(
            "UPDATE cap_workflow_instances SET current_step_key=%s WHERE workflow_instance_id=%s",
            (first_step["step_key"], instance_id),
        )
    cursor.execute(
        """
        INSERT INTO cap_workflow_tasks
          (workflow_instance_id, workflow_step_id, task_code, task_name, assigned_user_id, task_status, due_at, tenant_id)
        VALUES (%s, %s, %s, %s, %s, 'pending', %s, %s)
        """,
        (
            instance_id,
            step_id,
            f"WFT-API-{uuid.uuid4().hex[:8].upper()}",
            f"{step_name}({template['template_name']})",
            actor_user_id,
            datetime.now(),
            _tenant_for_user(cursor, actor_user_id),
        ),
    )
    task_id = int(cursor.lastrowid)
    audit_id = write_audit(
        cursor,
        actor_user_id,
        "workflow.start",
        "workflow_instance",
        instance_id,
        payload.title,
        after={"workflow_family": family, "task_id": task_id, "payload": payload.payload or {}},
    )
    return {
        "ok": True,
        "workflow_instance_id": instance_id,
        "task_id": task_id,
        "audit_id": audit_id,
        "affected_table": "cap_workflow_instances",
    }


@app.get("/health")
def health() -> dict[str, Any]:
    return {
        "ok": True,
        "app": APP_NAME,
        "version": APP_VERSION,
        "python": "3.12-required",
    }


@app.post("/api/auth/login")
def login(payload: LoginPayload) -> dict[str, Any]:
    account = payload.account.strip()
    if not account or not payload.password:
        raise HTTPException(status_code=400, detail="Account and password are required")

    connection = connect_db()
    try:
        with connection.cursor() as cursor:
            cursor.execute(
                """
                SELECT user_id, org_id, login_name, display_name, password_hash, account_status
                FROM cap_users WHERE login_name=%s AND deleted_at IS NULL LIMIT 1
                """,
                (account,),
            )
            user = cursor.fetchone()

            # 统一失败:不泄露账号是否存在;失败也记审计。
            if user is None or not verify_password(payload.password, user.get("password_hash")):
                cursor.execute(
                    """
                    INSERT INTO cap_login_events
                      (user_id, login_name, auth_method, outcome, ip_mask, device_label, risk_level, occurred_at)
                    VALUES (%s, %s, 'password', 'failed', 'api-client', 'deploy frontend', 'medium', %s)
                    """,
                    (user["user_id"] if user else None, account, datetime.now()),
                )
                connection.commit()
                raise HTTPException(status_code=401, detail="Invalid account or password")

            if user["account_status"] != "active":
                raise HTTPException(status_code=403, detail="Account is not active")

            # 真实角色来自 cap_user_roles(不再硬编码 managing_partner)。
            cursor.execute(
                """
                SELECT r.role_code FROM cap_user_roles ur
                JOIN cap_roles r ON r.role_id = ur.role_id
                WHERE ur.user_id = %s
                """,
                (user["user_id"],),
            )
            roles = [row["role_code"] for row in cursor.fetchall()]

            # 角色授予的权限码(RBAC):role → cap_role_permissions(effect=allow)。
            cursor.execute(
                """
                SELECT DISTINCT p.permission_code
                FROM cap_user_roles ur
                JOIN cap_role_permissions rp ON rp.role_id = ur.role_id AND rp.effect = 'allow'
                JOIN cap_permissions p ON p.permission_id = rp.permission_id
                WHERE ur.user_id = %s
                """,
                (user["user_id"],),
            )
            perms = [row["permission_code"] for row in cursor.fetchall()]

            # 解析租户 = 用户所属组织向上走到根公司(parent_org_id 为 NULL 者)。
            tenant_id = resolve_tenant_id(cursor, user["org_id"])

            cursor.execute(
                """
                INSERT INTO cap_login_events
                  (user_id, login_name, auth_method, outcome, ip_mask, device_label, risk_level, occurred_at)
                VALUES (%s, %s, 'password', 'success', 'api-client', 'deploy frontend', 'low', %s)
                """,
                (user["user_id"], account, datetime.now()),
            )
            cursor.execute("UPDATE cap_users SET last_login_at=%s WHERE user_id=%s", (datetime.now(), user["user_id"]))
            audit_id = write_audit(
                cursor,
                int(user["user_id"]),
                "auth.login",
                "user",
                int(user["user_id"]),
                str(user["display_name"]),
                after={"account": account, "outcome": "success"},
            )
        connection.commit()
    finally:
        connection.close()

    token = create_token(
        user_id=int(user["user_id"]),
        org_id=user["org_id"],
        tenant_id=tenant_id,
        roles=roles,
        perms=perms,
        display_name=str(user["display_name"]),
    )
    return {
        "token": token,
        "token_type": "bearer",
        "audit_id": audit_id,
        "user": {
            "id": user["user_id"],
            "account": account,
            "display_name": user["display_name"],
            "org_id": user["org_id"],
            "roles": roles,
            "perms": perms,
        },
    }


@app.get("/api/screens")
def screens() -> dict[str, Any]:
    connection = connect_db()
    try:
        with connection.cursor() as cursor:
            cursor.execute(
                """
                SELECT screen_code AS id, item_name AS title, group_name AS `group`, route_key
                FROM cap_navigation_items
                WHERE is_visible=1
                ORDER BY sort_order
                """
            )
            rows = cursor.fetchall()
    finally:
        connection.close()
    return {"count": len(rows), "items": rows}


@app.get("/api/dashboard")
def dashboard(user: AuthedUser = Depends(current_user)) -> dict[str, Any]:
    tid = tenant_of(user)
    connection = connect_db()
    try:
        with connection.cursor() as cursor:
            cursor.execute("SELECT COUNT(*) AS count FROM cap_projects WHERE deleted_at IS NULL AND tenant_id=%s", (tid,))
            project_count = cursor.fetchone()["count"]
            cursor.execute("SELECT COALESCE(SUM(committed_size),0) AS aum FROM cap_funds WHERE deleted_at IS NULL AND tenant_id=%s", (tid,))
            fund_aum = cursor.fetchone()["aum"]
            cursor.execute("SELECT COUNT(*) AS count FROM cap_workflow_tasks WHERE task_status='pending' AND tenant_id=%s", (tid,))
            pending_tasks = cursor.fetchone()["count"]
            cursor.execute("SELECT COUNT(*) AS count FROM cap_risk_incidents WHERE severity IN ('high','critical') AND deleted_at IS NULL AND tenant_id=%s", (tid,))
            high_risk = cursor.fetchone()["count"]
    finally:
        connection.close()

    return {
        "metrics": {
            "fund_aum": float(fund_aum),
            "project_reserve": project_count,
            "pending_tasks": pending_tasks,
            "high_risk": high_risk,
        }
    }


@app.get("/api/projects")
def list_projects(user: AuthedUser = Depends(current_user)) -> dict[str, Any]:
    connection = connect_db()
    try:
        with connection.cursor() as cursor:
            cursor.execute(
                """
                SELECT p.project_id AS id, p.short_name AS name, p.stage_label AS stage,
                       p.industry_group AS sector, p.city, u.display_name AS owner,
                       p.opportunity_status AS status, p.summary, p.updated_at
                FROM cap_projects p
                LEFT JOIN cap_users u ON u.user_id=p.owner_user_id
                WHERE p.deleted_at IS NULL AND p.tenant_id=%s
                ORDER BY p.project_id DESC
                """,
                (tenant_of(user),),
            )
            rows = cursor.fetchall()
    finally:
        connection.close()
    return {"count": len(rows), "items": rows}


@app.post("/api/projects")
def create_project(payload: CreateProjectPayload, user: AuthedUser = Depends(require_permission("project.edit"))) -> dict[str, Any]:
    connection = connect_db()
    try:
        with connection.cursor() as cursor:
            cursor.execute(
                """
                INSERT INTO cap_projects
                  (project_code, short_name, legal_name, opportunity_status, stage_label, industry_group,
                   city, registered_location, owner_user_id, source_channel, summary, thesis, product_note,
                   highlight_note, created_by, tenant_id)
                VALUES (%s, %s, %s, 'sourced', 'Sourced', %s, %s, %s, %s, 'API Create',
                        %s, 'API-created investment thesis pending review.', 'Product note pending.',
                        'Created from deploy backend.', %s, %s)
                """,
                (
                    f"PRJ-API-{uuid.uuid4().hex[:8].upper()}",
                    payload.short_name,
                    payload.legal_name or payload.short_name,
                    payload.industry_group,
                    payload.city,
                    payload.city,
                    user.user_id,
                    payload.summary or "Created through real backend API.",
                    user.user_id,
                    tenant_of(user),
                ),
            )
            project_id = int(cursor.lastrowid)
            cursor.execute(
                """
                INSERT INTO cap_project_stage_events
                  (project_id, from_stage, to_stage, event_reason, event_at, actor_user_id, notes, tenant_id)
                VALUES (%s, NULL, 'Sourced', 'Created through API', %s, %s, 'Initial API create', %s)
                """,
                (project_id, datetime.now(), user.user_id, tenant_of(user)),
            )
            audit_id = write_audit(
                cursor,
                user.user_id,
                "project.create",
                "project",
                project_id,
                payload.short_name,
                after=payload.model_dump(),
            )
        connection.commit()
    finally:
        connection.close()
    return {"ok": True, "project_id": project_id, "audit_id": audit_id}


@app.post("/api/projects/{project_id}/advance")
def advance_project(project_id: int, user: AuthedUser = Depends(require_permission("project.edit"))) -> dict[str, Any]:
    status_flow = [
        ("sourced", "screening", "Screening"),
        ("screening", "approved", "Project Approval"),
        ("approved", "term_sheet", "TS"),
        ("term_sheet", "diligence", "Diligence"),
        ("diligence", "committee", "IC"),
        ("committee", "agreement", "Agreement"),
        ("agreement", "funded", "Funded"),
        ("funded", "portfolio", "Post Investment"),
    ]
    connection = connect_db()
    try:
        with connection.cursor() as cursor:
            cursor.execute(
                "SELECT project_id, short_name, opportunity_status, stage_label FROM cap_projects WHERE project_id=%s AND tenant_id=%s",
                (project_id, tenant_of(user)),
            )
            row = cursor.fetchone()
            if row is None:
                raise HTTPException(status_code=404, detail="Project not found")

            next_status = None
            next_stage = None
            for current_status, candidate_status, candidate_stage in status_flow:
                if row["opportunity_status"] == current_status:
                    next_status = candidate_status
                    next_stage = candidate_stage
                    break
            if next_status is None:
                next_status = row["opportunity_status"]
                next_stage = row["stage_label"]

            cursor.execute(
                "UPDATE cap_projects SET opportunity_status=%s, stage_label=%s WHERE project_id=%s",
                (next_status, next_stage, project_id),
            )
            cursor.execute(
                """
                INSERT INTO cap_project_stage_events
                  (project_id, from_stage, to_stage, event_reason, event_at, actor_user_id, notes, tenant_id)
                VALUES (%s, %s, %s, 'Advanced through API', %s, %s, 'Stage advanced from frontend action', %s)
                """,
                (project_id, row["stage_label"], next_stage, datetime.now(), user.user_id, tenant_of(user)),
            )
            audit_id = write_audit(
                cursor,
                user.user_id,
                "project.advance_stage",
                "project",
                project_id,
                row["short_name"],
                before={"stage": row["stage_label"]},
                after={"stage": next_stage},
            )
        connection.commit()
    finally:
        connection.close()
    return {"ok": True, "project_id": project_id, "stage": next_stage, "audit_id": audit_id}


@app.get("/api/projects/{project_id}")
def get_project(project_id: int, user: AuthedUser = Depends(current_user)) -> dict[str, Any]:
    """单个项目详情(租户内)—— detail 屏加载真实实体数据。"""
    connection = connect_db()
    try:
        with connection.cursor() as cursor:
            cursor.execute(
                """
                SELECT p.project_id AS id, p.short_name, p.legal_name, p.opportunity_status,
                       p.stage_label, p.industry_group, p.city, p.registered_location,
                       p.registry_code_mask, p.source_channel, p.summary, p.thesis,
                       p.product_note, p.highlight_note, p.created_at,
                       u.display_name AS owner, p.updated_at
                FROM cap_projects p
                LEFT JOIN cap_users u ON u.user_id=p.owner_user_id
                WHERE p.project_id=%s AND p.tenant_id=%s AND p.deleted_at IS NULL
                """,
                (project_id, tenant_of(user)),
            )
            row = cursor.fetchone()
    finally:
        connection.close()
    if row is None:
        raise HTTPException(status_code=404, detail="Project not found")
    return row


@app.get("/api/projects/{project_id}/summary")
def project_investment_summary(project_id: int, user: AuthedUser = Depends(current_user)) -> dict[str, Any]:
    """投资汇总(项目卡片二级 tab):从 cap_investment_positions 聚合业绩/投资/回款指标。
    MOIC/DPI 由真实金额算出;IRR 需现金流序列,暂不估(返回 None)。"""
    tid = tenant_of(user)
    connection = connect_db()
    try:
        with connection.cursor() as cursor:
            cursor.execute(
                """
                SELECT
                  SUM(agreement_amount)        AS agreement_total,
                  SUM(cumulative_paid_amount)  AS paid_total,
                  SUM(realized_return_amount)  AS realized_total,
                  MAX(latest_valuation)        AS latest_valuation,
                  MAX(current_ownership_ratio) AS ownership_ratio,
                  MIN(first_payment_on)        AS first_payment_on,
                  COUNT(*)                     AS position_count,
                  MAX(round_label)             AS round_label,
                  MAX(investment_status)       AS investment_status,
                  MAX(exit_status)             AS exit_status
                FROM cap_investment_positions
                WHERE project_id=%s AND tenant_id=%s AND deleted_at IS NULL
                """,
                (project_id, tid),
            )
            agg = cursor.fetchone() or {}
    finally:
        connection.close()

    def fnum(v: Any) -> float | None:
        return float(v) if v is not None else None

    paid = fnum(agg.get("paid_total")) or 0.0
    valuation = fnum(agg.get("latest_valuation"))
    realized = fnum(agg.get("realized_total")) or 0.0
    # MOIC=(当前持有估值+已实现回款)/已投入;DPI=已实现回款/已投入。paid=0 时不计算(避免除零)。
    moic = round((((valuation or 0.0) + realized) / paid), 2) if paid > 0 else None
    dpi = round((realized / paid), 2) if paid > 0 else None
    return {
        "has_position": (agg.get("position_count") or 0) > 0,
        "performance": {"DPI": dpi, "MOIC": moic, "IRR": None},
        "investment": {
            "agreement_total": fnum(agg.get("agreement_total")),
            "paid_total": fnum(agg.get("paid_total")),
            "first_payment_on": agg.get("first_payment_on").isoformat() if agg.get("first_payment_on") else None,
            "ownership_ratio": fnum(agg.get("ownership_ratio")),
            "latest_valuation": valuation,
            "round_label": agg.get("round_label"),
            "investment_status": agg.get("investment_status"),
        },
        "realized": {"realized_total": realized or None, "exit_status": agg.get("exit_status")},
    }


@app.patch("/api/projects/{project_id}")
def update_project(
    project_id: int,
    payload: UpdateProjectPayload,
    user: AuthedUser = Depends(require_permission("project.edit")),
) -> dict[str, Any]:
    """编辑项目字段(租户内 + project.edit)—— detail 屏保存。只更新传入的字段。"""
    fields = {k: v for k, v in payload.model_dump().items() if v is not None}
    expected = fields.pop("expected_updated_at", None)  # 乐观锁字段不入 SET
    if not fields:
        raise HTTPException(status_code=400, detail="No fields to update")
    connection = connect_db()
    try:
        with connection.cursor() as cursor:
            cursor.execute(
                "SELECT short_name, updated_at FROM cap_projects WHERE project_id=%s AND tenant_id=%s AND deleted_at IS NULL",
                (project_id, tenant_of(user)),
            )
            before = cursor.fetchone()
            if before is None:
                raise HTTPException(status_code=404, detail="Project not found")
            # 乐观锁:加载后若被他人改过,updated_at 会变,拒绝覆盖。
            if expected and before["updated_at"] and before["updated_at"].isoformat() != expected:
                raise HTTPException(status_code=409, detail="记录已被他人修改,请刷新后重试")
            set_clause = ", ".join(f"{col}=%s" for col in fields)
            cursor.execute(
                f"UPDATE cap_projects SET {set_clause} WHERE project_id=%s AND tenant_id=%s",
                (*fields.values(), project_id, tenant_of(user)),
            )
            audit_id = write_audit(
                cursor, user.user_id, "project.update", "project", project_id,
                fields.get("short_name", before["short_name"]), after=fields,
            )
        connection.commit()
    finally:
        connection.close()
    return {"ok": True, "project_id": project_id, "updated_fields": list(fields), "audit_id": audit_id}


def _soft_delete(cursor, *, table: str, id_col: str, obj_type: str, obj_id: int, label: str, user: "AuthedUser") -> None:
    """软删除:置 deleted_at + 写回收站行(可 30 天内恢复)。"""
    now = datetime.now()
    tid = tenant_of(user)
    cursor.execute(f"UPDATE {table} SET deleted_at=%s WHERE {id_col}=%s AND tenant_id=%s", (now, obj_id, tid))
    cursor.execute(
        """
        INSERT INTO cap_recycle_items
          (object_type, object_id, object_label, purge_status, deleted_by, deleted_at, recoverable_until, tenant_id)
        VALUES (%s, %s, %s, 'recoverable', %s, %s, %s, %s)
        """,
        (obj_type, obj_id, label, user.user_id, now, now + timedelta(days=30), tid),
    )


@app.delete("/api/projects/{project_id}")
def delete_project(project_id: int, user: AuthedUser = Depends(require_permission("project.edit"))) -> dict[str, Any]:
    """软删除项目 → 进回收站(租户内 + project.edit)。"""
    connection = connect_db()
    try:
        with connection.cursor() as cursor:
            cursor.execute(
                "SELECT short_name FROM cap_projects WHERE project_id=%s AND tenant_id=%s AND deleted_at IS NULL",
                (project_id, tenant_of(user)),
            )
            row = cursor.fetchone()
            if row is None:
                raise HTTPException(status_code=404, detail="Project not found")
            _soft_delete(cursor, table="cap_projects", id_col="project_id", obj_type="project", obj_id=project_id, label=row["short_name"], user=user)
            audit_id = write_audit(cursor, user.user_id, "project.delete", "project", project_id, row["short_name"], risk_level="high")
        connection.commit()
    finally:
        connection.close()
    return {"ok": True, "project_id": project_id, "audit_id": audit_id}


@app.get("/api/funds")
def list_funds(user: AuthedUser = Depends(current_user)) -> dict[str, Any]:
    connection = connect_db()
    try:
        with connection.cursor() as cursor:
            cursor.execute(
                """
                SELECT fund_id AS id, fund_name AS name, fund_status AS status,
                       committed_size, paid_in_size, net_asset_value, unit_nav
                FROM cap_funds
                WHERE deleted_at IS NULL AND tenant_id=%s
                ORDER BY fund_id
                """,
                (tenant_of(user),),
            )
            rows = cursor.fetchall()
    finally:
        connection.close()
    return {"count": len(rows), "items": rows}


@app.post("/api/funds")
def create_fund(payload: CreateFundPayload, user: AuthedUser = Depends(require_roles("system_admin", "managing_partner", "fund_operator"))) -> dict[str, Any]:
    connection = connect_db()
    try:
        with connection.cursor() as cursor:
            result = create_fund_record(cursor, payload, user.user_id)
        connection.commit()
    finally:
        connection.close()
    return result


@app.get("/api/funds/{fund_id}")
def get_fund(fund_id: int, user: AuthedUser = Depends(current_user)) -> dict[str, Any]:
    """单个基金详情(租户内)。"""
    connection = connect_db()
    try:
        with connection.cursor() as cursor:
            cursor.execute(
                """
                SELECT fund_id AS id, fund_name, legal_name, fund_status, raise_method,
                       target_size, committed_size, paid_in_size, net_asset_value, unit_nav, updated_at
                FROM cap_funds
                WHERE fund_id=%s AND tenant_id=%s AND deleted_at IS NULL
                """,
                (fund_id, tenant_of(user)),
            )
            row = cursor.fetchone()
    finally:
        connection.close()
    if row is None:
        raise HTTPException(status_code=404, detail="Fund not found")
    return row


@app.patch("/api/funds/{fund_id}")
def update_fund(
    fund_id: int,
    payload: UpdateFundPayload,
    user: AuthedUser = Depends(require_roles("system_admin", "managing_partner", "fund_operator")),
) -> dict[str, Any]:
    """编辑基金字段(租户内 + 基金运营类角色)。"""
    fields = {k: v for k, v in payload.model_dump().items() if v is not None}
    expected = fields.pop("expected_updated_at", None)
    if not fields:
        raise HTTPException(status_code=400, detail="No fields to update")
    connection = connect_db()
    try:
        with connection.cursor() as cursor:
            cursor.execute(
                "SELECT fund_name, updated_at FROM cap_funds WHERE fund_id=%s AND tenant_id=%s AND deleted_at IS NULL",
                (fund_id, tenant_of(user)),
            )
            before = cursor.fetchone()
            if before is None:
                raise HTTPException(status_code=404, detail="Fund not found")
            if expected and before["updated_at"] and before["updated_at"].isoformat() != expected:
                raise HTTPException(status_code=409, detail="记录已被他人修改,请刷新后重试")
            set_clause = ", ".join(f"{col}=%s" for col in fields)
            cursor.execute(
                f"UPDATE cap_funds SET {set_clause} WHERE fund_id=%s AND tenant_id=%s",
                (*fields.values(), fund_id, tenant_of(user)),
            )
            audit_id = write_audit(
                cursor, user.user_id, "fund.update", "fund", fund_id,
                fields.get("fund_name", before["fund_name"]), after=fields,
            )
        connection.commit()
    finally:
        connection.close()
    return {"ok": True, "fund_id": fund_id, "updated_fields": list(fields), "audit_id": audit_id}


@app.delete("/api/funds/{fund_id}")
def delete_fund(fund_id: int, user: AuthedUser = Depends(require_roles("system_admin", "managing_partner", "fund_operator"))) -> dict[str, Any]:
    """软删除基金 → 进回收站。"""
    connection = connect_db()
    try:
        with connection.cursor() as cursor:
            cursor.execute(
                "SELECT fund_name FROM cap_funds WHERE fund_id=%s AND tenant_id=%s AND deleted_at IS NULL",
                (fund_id, tenant_of(user)),
            )
            row = cursor.fetchone()
            if row is None:
                raise HTTPException(status_code=404, detail="Fund not found")
            _soft_delete(cursor, table="cap_funds", id_col="fund_id", obj_type="fund", obj_id=fund_id, label=row["fund_name"], user=user)
            audit_id = write_audit(cursor, user.user_id, "fund.delete", "fund", fund_id, row["fund_name"], risk_level="high")
        connection.commit()
    finally:
        connection.close()
    return {"ok": True, "fund_id": fund_id, "audit_id": audit_id}


@app.get("/api/risks")
def list_risks(user: AuthedUser = Depends(current_user)) -> dict[str, Any]:
    connection = connect_db()
    try:
        with connection.cursor() as cursor:
            cursor.execute(
                """
                SELECT r.risk_incident_id AS id, r.incident_title AS title, r.severity,
                       r.incident_status AS status, r.latest_progress, p.short_name AS project
                FROM cap_risk_incidents r
                LEFT JOIN cap_projects p ON p.project_id=r.project_id
                WHERE r.deleted_at IS NULL AND r.tenant_id=%s
                ORDER BY r.risk_incident_id DESC
                """,
                (tenant_of(user),),
            )
            rows = cursor.fetchall()
    finally:
        connection.close()
    return {"count": len(rows), "items": rows}


@app.post("/api/risks/{risk_incident_id}/updates")
def add_risk_update(
    risk_incident_id: int,
    payload: RiskUpdatePayload,
    user: AuthedUser = Depends(require_permission("risk.manage")),
) -> dict[str, Any]:
    connection = connect_db()
    try:
        with connection.cursor() as cursor:
            cursor.execute(
                "SELECT incident_title FROM cap_risk_incidents WHERE risk_incident_id=%s AND tenant_id=%s",
                (risk_incident_id, tenant_of(user)),
            )
            row = cursor.fetchone()
            if row is None:
                raise HTTPException(status_code=404, detail="Risk incident not found")
            cursor.execute(
                """
                INSERT INTO cap_risk_updates (risk_incident_id, update_text, update_status, created_by, tenant_id)
                VALUES (%s, %s, %s, %s, %s)
                """,
                (risk_incident_id, payload.update_text, payload.update_status, user.user_id, tenant_of(user)),
            )
            update_id = int(cursor.lastrowid)
            cursor.execute(
                "UPDATE cap_risk_incidents SET latest_progress=%s, incident_status='mitigating' WHERE risk_incident_id=%s AND tenant_id=%s",
                (payload.update_text, risk_incident_id, tenant_of(user)),
            )
            audit_id = write_audit(
                cursor,
                user.user_id,
                "risk.update",
                "risk_incident",
                risk_incident_id,
                row["incident_title"],
                after=payload.model_dump(),
                risk_level="medium",
            )
        connection.commit()
    finally:
        connection.close()
    return {"ok": True, "update_id": update_id, "audit_id": audit_id}


@app.get("/api/workflow/tasks")
def workflow_tasks(user: AuthedUser = Depends(current_user)) -> dict[str, Any]:
    connection = connect_db()
    try:
        with connection.cursor() as cursor:
            cursor.execute(
                """
                SELECT t.workflow_task_id AS id, t.task_name, t.task_status, t.due_at,
                       i.title AS instance_title, i.instance_status, u.display_name AS assignee
                FROM cap_workflow_tasks t
                JOIN cap_workflow_instances i ON i.workflow_instance_id=t.workflow_instance_id
                LEFT JOIN cap_users u ON u.user_id=t.assigned_user_id
                WHERE t.tenant_id=%s
                ORDER BY t.workflow_task_id DESC
                """,
                (tenant_of(user),),
            )
            rows = cursor.fetchall()
    finally:
        connection.close()
    return {"count": len(rows), "items": rows}


@app.get("/api/workflow/templates")
def workflow_templates(user: AuthedUser = Depends(current_user)) -> dict[str, Any]:
    """流程模板 + 有序步骤(前端据此渲染真实编排链,替代写死的模板数)。"""
    connection = connect_db()
    try:
        with connection.cursor() as cursor:
            cursor.execute(
                "SELECT workflow_template_id AS id, workflow_family, template_name FROM cap_workflow_templates ORDER BY workflow_family, workflow_template_id"
            )
            templates = cursor.fetchall()
            cursor.execute(
                "SELECT workflow_template_id AS tid, step_key, step_name, step_type, sort_order FROM cap_workflow_steps ORDER BY workflow_template_id, sort_order"
            )
            steps = cursor.fetchall()
    finally:
        connection.close()
    by_tpl: dict[int, list[dict[str, Any]]] = {}
    for s in steps:
        by_tpl.setdefault(s["tid"], []).append(s)
    for t in templates:
        t["steps"] = by_tpl.get(t["id"], [])
    return {"count": len(templates), "items": templates}


@app.post("/api/workflow/instances")
def start_workflow(payload: StartWorkflowPayload, user: AuthedUser = Depends(require_roles("system_admin", "managing_partner", "fund_operator", "risk_legal", "investment_manager"))) -> dict[str, Any]:
    connection = connect_db()
    try:
        with connection.cursor() as cursor:
            result = start_workflow_record(cursor, payload, user.user_id)
        connection.commit()
    finally:
        connection.close()
    return result


@app.post("/api/workflow/tasks/{task_id}/action")
def act_workflow_task(
    task_id: int,
    payload: WorkflowActionPayload,
    user: AuthedUser = Depends(require_roles("system_admin", "managing_partner", "fund_operator", "risk_legal", "investment_manager")),
) -> dict[str, Any]:
    allowed = {"approve": "approved", "reject": "rejected", "transfer": "transferred", "archive": "archived"}
    if payload.action not in allowed:
        raise HTTPException(status_code=400, detail="Unsupported workflow action")

    tid = tenant_of(user)
    now = datetime.now()
    connection = connect_db()
    try:
        with connection.cursor() as cursor:
            # 连 instance/step/template 一起取,才能做流转。
            cursor.execute(
                """
                SELECT t.task_name, t.task_status, t.workflow_instance_id, t.workflow_step_id,
                       i.title AS instance_title, i.instance_status,
                       i.workflow_template_id, s.sort_order AS cur_order
                FROM cap_workflow_tasks t
                JOIN cap_workflow_instances i ON i.workflow_instance_id=t.workflow_instance_id
                LEFT JOIN cap_workflow_steps s ON s.workflow_step_id=t.workflow_step_id
                WHERE t.workflow_task_id=%s AND t.tenant_id=%s
                """,
                (task_id, tid),
            )
            row = cursor.fetchone()
            if row is None:
                raise HTTPException(status_code=404, detail="Workflow task not found")
            if row["task_status"] != "pending":
                raise HTTPException(status_code=409, detail=f"任务已处理({row['task_status']}),不能重复操作")

            next_status = allowed[payload.action]
            cursor.execute(
                "UPDATE cap_workflow_tasks SET task_status=%s, acted_at=%s, action_comment=%s WHERE workflow_task_id=%s AND tenant_id=%s",
                (next_status, now, payload.comment, task_id, tid),
            )

            # 流转:approve → 有下一步则建下一步任务并推进实例;无则实例 approved(闭环)。
            #       reject → 实例 rejected(终止,退回发起人)。transfer/archive 只作用于该任务。
            instance_id = row["workflow_instance_id"]
            flow: dict[str, Any] = {}
            if payload.action == "approve":
                next_step = None
                if row["cur_order"] is not None:
                    cursor.execute(
                        """
                        SELECT workflow_step_id, step_key, step_name FROM cap_workflow_steps
                        WHERE workflow_template_id=%s AND sort_order>%s
                        ORDER BY sort_order, workflow_step_id LIMIT 1
                        """,
                        (row["workflow_template_id"], row["cur_order"]),
                    )
                    next_step = cursor.fetchone()
                if next_step:
                    cursor.execute(
                        """
                        INSERT INTO cap_workflow_tasks
                          (workflow_instance_id, workflow_step_id, task_code, task_name, assigned_user_id, task_status, due_at, tenant_id)
                        VALUES (%s, %s, %s, %s, %s, 'pending', %s, %s)
                        """,
                        (instance_id, next_step["workflow_step_id"], f"WFT-API-{uuid.uuid4().hex[:8].upper()}",
                         next_step["step_name"], user.user_id, now, tid),
                    )
                    new_task_id = int(cursor.lastrowid)
                    cursor.execute(
                        "UPDATE cap_workflow_instances SET current_step_key=%s WHERE workflow_instance_id=%s AND tenant_id=%s",
                        (next_step["step_key"], instance_id, tid),
                    )
                    flow = {"instance_status": "running", "next_step": next_step["step_name"], "next_task_id": new_task_id}
                else:
                    cursor.execute(
                        "UPDATE cap_workflow_instances SET instance_status='approved', current_step_key='end', completed_at=%s WHERE workflow_instance_id=%s AND tenant_id=%s",
                        (now, instance_id, tid),
                    )
                    flow = {"instance_status": "approved", "next_step": None}
            elif payload.action == "reject":
                cursor.execute(
                    "UPDATE cap_workflow_instances SET instance_status='rejected', completed_at=%s WHERE workflow_instance_id=%s AND tenant_id=%s",
                    (now, instance_id, tid),
                )
                flow = {"instance_status": "rejected", "next_step": None}
            else:
                flow = {"instance_status": row["instance_status"], "next_step": None}

            audit_id = write_audit(
                cursor,
                user.user_id,
                f"workflow.{payload.action}",
                "workflow_task",
                task_id,
                row["task_name"],
                before={"status": row["task_status"]},
                after={"status": next_status, "comment": payload.comment, **flow},
                risk_level="medium" if payload.action == "reject" else "low",
            )
        connection.commit()
    finally:
        connection.close()
    return {"ok": True, "task_id": task_id, "status": next_status, "audit_id": audit_id, **flow}


@app.get("/api/documents")
def list_documents(user: AuthedUser = Depends(current_user)) -> dict[str, Any]:
    connection = connect_db()
    try:
        with connection.cursor() as cursor:
            cursor.execute(
                """
                SELECT document_id AS id, title, document_kind, file_name, access_level,
                       watermark_policy, current_version_no, fulltext_status, storage_uri
                FROM cap_documents
                WHERE deleted_at IS NULL AND tenant_id=%s
                ORDER BY document_id DESC
                """,
                (tenant_of(user),),
            )
            rows = cursor.fetchall()
    finally:
        connection.close()
    return {"count": len(rows), "items": rows}


@app.post("/api/documents")
def upload_document(payload: CreateDocumentPayload, user: AuthedUser = Depends(require_roles("system_admin", "managing_partner", "investment_manager", "risk_legal"))) -> dict[str, Any]:
    connection = connect_db()
    try:
        with connection.cursor() as cursor:
            result = create_document_record(cursor, payload, user.user_id)
        connection.commit()
    finally:
        connection.close()
    return result


@app.post("/api/documents/{document_id}/download")
def download_document(document_id: int, user: AuthedUser = Depends(require_permission("document.download"))) -> dict[str, Any]:
    connection = connect_db()
    try:
        with connection.cursor() as cursor:
            cursor.execute(
                "SELECT title, storage_uri FROM cap_documents WHERE document_id=%s AND tenant_id=%s",
                (document_id, tenant_of(user)),
            )
            row = cursor.fetchone()
            if row is None:
                raise HTTPException(status_code=404, detail="Document not found")
            audit_id = write_audit(
                cursor,
                user.user_id,
                "document.download",
                "document",
                document_id,
                row["title"],
                after={"secondary_auth": True, "storage_uri": row["storage_uri"]},
                risk_level="high",
            )
        connection.commit()
    finally:
        connection.close()
    return {"ok": True, "document_id": document_id, "download_uri": row["storage_uri"], "audit_id": audit_id}


@app.post("/api/import-export/tasks")
def create_import_export_task(payload: ImportExportPayload, user: AuthedUser = Depends(require_permission("report.export", "fund.export"))) -> dict[str, Any]:
    connection = connect_db()
    try:
        with connection.cursor() as cursor:
            cursor.execute(
                """
                INSERT INTO cap_import_export_tasks
                  (task_code, task_kind, entity_type, task_status, source_file_uri, result_file_uri,
                   total_rows, success_rows, failed_rows, error_summary, requested_by, requested_at, tenant_id)
                VALUES (%s, %s, %s, 'queued', %s, NULL, 0, 0, 0, NULL, %s, %s, %s)
                """,
                (
                    f"IET-API-{uuid.uuid4().hex[:8].upper()}",
                    payload.task_kind,
                    payload.entity_type,
                    payload.source_file_uri,
                    user.user_id,
                    datetime.now(),
                    tenant_of(user),
                ),
            )
            task_id = int(cursor.lastrowid)
            audit_id = write_audit(
                cursor,
                user.user_id,
                f"import_export.{payload.task_kind}",
                "import_export_task",
                task_id,
                payload.entity_type,
                after=payload.model_dump(),
            )
        connection.commit()
    finally:
        connection.close()
    return {"ok": True, "task_id": task_id, "audit_id": audit_id}


_RECYCLE_TABLES = {"project": ("cap_projects", "project_id"), "fund": ("cap_funds", "fund_id")}


@app.post("/api/recycle/{recycle_item_id}/restore")
def restore_recycle_item(recycle_item_id: int, user: AuthedUser = Depends(require_roles("system_admin", "managing_partner"))) -> dict[str, Any]:
    """恢复回收站条目:标记 restored 并把对应实体 deleted_at 置空(真回滚软删除)。"""
    connection = connect_db()
    try:
        with connection.cursor() as cursor:
            cursor.execute(
                "SELECT object_label, object_type, object_id, purge_status FROM cap_recycle_items WHERE recycle_item_id=%s AND tenant_id=%s",
                (recycle_item_id, tenant_of(user)),
            )
            row = cursor.fetchone()
            if row is None:
                raise HTTPException(status_code=404, detail="Recycle item not found")
            if row["purge_status"] != "recoverable":
                raise HTTPException(status_code=409, detail="该条目不可恢复(已恢复或已彻底删除)")
            cursor.execute(
                """
                UPDATE cap_recycle_items
                SET purge_status='restored', restored_by=%s, restored_at=%s
                WHERE recycle_item_id=%s AND tenant_id=%s
                """,
                (user.user_id, datetime.now(), recycle_item_id, tenant_of(user)),
            )
            tbl = _RECYCLE_TABLES.get(row["object_type"])
            if tbl:
                cursor.execute(
                    f"UPDATE {tbl[0]} SET deleted_at=NULL WHERE {tbl[1]}=%s AND tenant_id=%s",
                    (row["object_id"], tenant_of(user)),
                )
            audit_id = write_audit(
                cursor,
                user.user_id,
                "recycle.restore",
                "recycle_item",
                recycle_item_id,
                row["object_label"],
                before={"status": row["purge_status"]},
                after={"status": "restored"},
                risk_level="medium",
            )
        connection.commit()
    finally:
        connection.close()
    return {"ok": True, "recycle_item_id": recycle_item_id, "audit_id": audit_id}


@app.post("/api/recycle/{recycle_item_id}/purge")
def purge_recycle_item(recycle_item_id: int, user: AuthedUser = Depends(require_roles("system_admin", "managing_partner"))) -> dict[str, Any]:
    """彻底删除:标记 purged(不可恢复)并硬删除底层已软删的实体,回收站列表不再显示。"""
    connection = connect_db()
    try:
        with connection.cursor() as cursor:
            cursor.execute(
                "SELECT object_label, object_type, object_id, purge_status FROM cap_recycle_items WHERE recycle_item_id=%s AND tenant_id=%s",
                (recycle_item_id, tenant_of(user)),
            )
            row = cursor.fetchone()
            if row is None:
                raise HTTPException(status_code=404, detail="Recycle item not found")
            cursor.execute(
                "UPDATE cap_recycle_items SET purge_status='purged' WHERE recycle_item_id=%s AND tenant_id=%s",
                (recycle_item_id, tenant_of(user)),
            )
            # 硬删已软删的实体(仅当处于软删状态,避免误删已恢复的)。
            tbl = _RECYCLE_TABLES.get(row["object_type"])
            if tbl:
                try:
                    cursor.execute(
                        f"DELETE FROM {tbl[0]} WHERE {tbl[1]}=%s AND tenant_id=%s AND deleted_at IS NOT NULL",
                        (row["object_id"], tenant_of(user)),
                    )
                except Exception:  # noqa: BLE001 —— 有外键引用则保留软删,仅标 purged
                    pass
            audit_id = write_audit(cursor, user.user_id, "recycle.purge", "recycle_item", recycle_item_id, row["object_label"], risk_level="high")
        connection.commit()
    finally:
        connection.close()
    return {"ok": True, "recycle_item_id": recycle_item_id, "audit_id": audit_id}


@app.post("/api/settings/preferences")
def save_preferences(payload: PreferencePayload, user: AuthedUser = Depends(current_user)) -> dict[str, Any]:
    connection = connect_db()
    try:
        with connection.cursor() as cursor:
            cursor.execute(
                """
                INSERT INTO cap_user_preferences
                  (user_id, notification_json, favorite_nav_json, table_view_json)
                VALUES (%s, %s, %s, %s)
                ON DUPLICATE KEY UPDATE
                  notification_json=VALUES(notification_json),
                  favorite_nav_json=VALUES(favorite_nav_json),
                  table_view_json=VALUES(table_view_json)
                """,
                (
                    user.user_id,
                    dumps(payload.notification_json or {}, ensure_ascii=False),
                    dumps(payload.favorite_nav_json or [], ensure_ascii=False),
                    dumps(payload.table_view_json or {}, ensure_ascii=False),
                ),
            )
            audit_id = write_audit(
                cursor,
                user.user_id,
                "settings.preferences.save",
                "user",
                user.user_id,
                "preferences",
                after=payload.model_dump(),
            )
        connection.commit()
    finally:
        connection.close()
    return {"ok": True, "audit_id": audit_id}


@app.get("/api/ledger/{screen_id}")
def screen_ledger(
    screen_id: str,
    user: AuthedUser = Depends(current_user),
    page: int = 1,
    page_size: int = 20,
    q: str = "",
) -> dict[str, Any]:
    query = LEDGER_QUERIES.get(screen_id)
    if query is None:
        raise HTTPException(status_code=404, detail=f"No ledger mapping for {screen_id}")

    page = max(1, page)
    page_size = min(200, max(1, page_size))  # 上限 200,防一次拉爆
    offset = (page - 1) * page_size
    tenant_params: tuple[Any, ...] = (tenant_of(user),) if screen_id in LEDGER_TENANT_SCOPED else ()
    keyword = q.strip().lower()

    connection = connect_db()
    try:
        with connection.cursor() as cursor:
            if keyword:
                # 高级筛选:服务端跨全部数据(非仅当前页)按关键词过滤。取上限 5000 行,
                # 在服务端逐值匹配(隐藏 __ 列不参与),再分页。
                cursor.execute(f"{query}\n            LIMIT 5000", tenant_params)
                all_rows = cursor.fetchall()
                matched = [
                    r for r in all_rows
                    if any(keyword in str(v).lower() for k, v in r.items() if not str(k).startswith("__") and v is not None)
                ]
                total = len(matched)
                rows = matched[offset:offset + page_size]
            else:
                base_no_order = query.split("ORDER BY")[0]
                cursor.execute(f"SELECT COUNT(*) AS n FROM ({base_no_order}) AS _c", tenant_params)
                total = int(cursor.fetchone()["n"])
                cursor.execute(f"{query}\n            LIMIT %s OFFSET %s", (*tenant_params, page_size, offset))
                rows = cursor.fetchall()
    finally:
        connection.close()
    return {
        "ok": True,
        "screen_id": screen_id,
        "source": "mysql",
        "count": len(rows),
        "total": total,
        "page": page,
        "page_size": page_size,
        "items": rows,
    }


@app.post("/api/screens/{screen_id}/primary-action")
def run_screen_primary_action(
    screen_id: str,
    user: AuthedUser = Depends(require_roles("system_admin", "managing_partner", "investment_manager", "fund_operator", "risk_legal")),
) -> dict[str, Any]:
    action_code = f"screen.{screen_id}.primary_action"
    connection = connect_db()
    try:
        with connection.cursor() as cursor:
            entity_type = "screen"
            entity_id: int | None = None
            entity_label = screen_id
            result: dict[str, Any]

            if screen_id in {"project-list", "project-board"}:
                payload = CreateProjectPayload(
                    short_name=f"Frontend Project {uuid.uuid4().hex[:4].upper()}",
                    legal_name="Frontend Project Co",
                    industry_group="Enterprise Software",
                    city="Shanghai",
                    summary="Created by screen primary action.",
                )
                cursor.execute(
                    """
                    INSERT INTO cap_projects
                      (project_code, short_name, legal_name, opportunity_status, stage_label,
                       industry_group, city, registered_location, owner_user_id, source_channel,
                       summary, thesis, product_note, highlight_note, created_by, tenant_id)
                    VALUES (%s, %s, %s, 'sourced', 'Sourced', %s, %s, %s, %s,
                            'Screen Primary Action', %s, 'Pending thesis.', 'Pending product note.',
                            'Created from page header.', %s, %s)
                    """,
                    (
                        f"PRJ-HDR-{uuid.uuid4().hex[:8].upper()}",
                        payload.short_name,
                        payload.legal_name,
                        payload.industry_group,
                        payload.city,
                        payload.city,
                        user.user_id,
                        payload.summary,
                        user.user_id,
                        tenant_of(user),
                    ),
                )
                entity_id = int(cursor.lastrowid)
                entity_type = "project"
                entity_label = payload.short_name
                result = {"ok": True, "project_id": entity_id, "affected_table": "cap_projects"}

            elif screen_id in {"fund-list", "fund-add"}:
                payload = CreateFundPayload(fund_name=f"Frontend Fund {uuid.uuid4().hex[:4].upper()}", target_size=100000000)
                result = create_fund_record(cursor, payload, user.user_id)
                entity_type = "fund"
                entity_id = int(result["fund_id"])
                entity_label = payload.fund_name

            elif screen_id == "announcements":
                entity_label = f"Frontend Announcement {uuid.uuid4().hex[:4].upper()}"
                cursor.execute(
                    """
                    INSERT INTO cap_announcements
                      (announcement_code, title, body_text, audience_scope_json, publish_status, published_at, created_by, tenant_id)
                    VALUES (%s, %s, 'Created by screen primary action.', %s, 'published', %s, %s, %s)
                    """,
                    (
                        f"ANN-HDR-{uuid.uuid4().hex[:8].upper()}",
                        entity_label,
                        dumps({"scope": "company"}, ensure_ascii=False),
                        datetime.now(),
                        user.user_id,
                        tenant_of(user),
                    ),
                )
                entity_type = "announcement"
                entity_id = int(cursor.lastrowid)
                result = {"ok": True, "announcement_id": entity_id, "affected_table": "cap_announcements"}

            elif screen_id == "calendar":
                entity_label = f"Frontend Calendar {uuid.uuid4().hex[:4].upper()}"
                starts_at = datetime.now()
                cursor.execute(
                    """
                    INSERT INTO cap_calendar_events
                      (event_code, event_title, event_kind, linked_entity_type, starts_at, ends_at,
                       organizer_user_id, location_text, visibility, tenant_id)
                    VALUES (%s, %s, 'project', 'project', %s, %s, %s, 'Online', 'department', %s)
                    """,
                    (
                        f"CAL-HDR-{uuid.uuid4().hex[:8].upper()}",
                        entity_label,
                        starts_at,
                        starts_at,
                        user.user_id,
                        tenant_of(user),
                    ),
                )
                entity_type = "calendar_event"
                entity_id = int(cursor.lastrowid)
                result = {"ok": True, "calendar_event_id": entity_id, "affected_table": "cap_calendar_events"}

            elif screen_id == "message-center":
                cursor.execute(
                    """
                    UPDATE cap_messages
                    SET read_at=COALESCE(read_at, %s), message_box='all', action_status='dismissed'
                    WHERE recipient_user_id=%s AND read_at IS NULL
                    """,
                    (datetime.now(), user.user_id),
                )
                affected = int(cursor.rowcount)
                entity_type = "message"
                entity_label = "mark_all_read"
                result = {"ok": True, "affected_rows": affected, "affected_table": "cap_messages"}

            elif screen_id == "investor-list":
                entity_label = f"Frontend LP {uuid.uuid4().hex[:4].upper()}"
                cursor.execute(
                    """
                    INSERT INTO cap_investors
                      (investor_code, investor_name, investor_kind, qualification_status, risk_rating,
                       city, country_code, disclosure_status, owner_user_id, created_by, tenant_id)
                    VALUES (%s, %s, 'institution', 'pending', 'professional', 'Shanghai', 'CN', 'not_started', %s, %s, %s)
                    """,
                    (f"INV-HDR-{uuid.uuid4().hex[:8].upper()}", entity_label, user.user_id, user.user_id, tenant_of(user)),
                )
                entity_type = "investor"
                entity_id = int(cursor.lastrowid)
                result = {"ok": True, "investor_id": entity_id, "affected_table": "cap_investors"}

            elif screen_id in {"flow-center", "flow-project", "flow-fund", "flow-oa"}:
                family = "fund" if screen_id == "flow-fund" else "office" if screen_id == "flow-oa" else "project"
                payload = StartWorkflowPayload(title=f"Frontend Workflow {uuid.uuid4().hex[:4].upper()}", workflow_family=family)
                result = start_workflow_record(cursor, payload, user.user_id)
                entity_type = "workflow_instance"
                entity_id = int(result["workflow_instance_id"])
                entity_label = payload.title

            elif screen_id in {"document-center", "process-files"}:
                kind = "workflow" if screen_id == "process-files" else "shared"
                payload = CreateDocumentPayload(title=f"Frontend Document {uuid.uuid4().hex[:4].upper()}", document_kind=kind)
                result = create_document_record(cursor, payload, user.user_id)
                entity_type = "document"
                entity_id = int(result["document_id"])
                entity_label = payload.title

            elif screen_id in {"ai-workspace", "meeting-ai"}:
                parse_kind = "meeting_minutes" if screen_id == "meeting-ai" else "business_plan"
                entity_label = f"Frontend AI Job {uuid.uuid4().hex[:4].upper()}"
                cursor.execute(
                    """
                    INSERT INTO cap_ai_parse_jobs
                      (job_code, job_name, parse_kind, requested_by, job_status)
                    VALUES (%s, %s, %s, %s, 'queued')
                    """,
                    (f"AI-HDR-{uuid.uuid4().hex[:8].upper()}", entity_label, parse_kind, user.user_id),
                )
                entity_type = "ai_parse_job"
                entity_id = int(cursor.lastrowid)
                result = {"ok": True, "ai_parse_job_id": entity_id, "affected_table": "cap_ai_parse_jobs"}

            elif screen_id == "system-users":
                suffix = uuid.uuid4().hex[:6]
                entity_label = f"Frontend User {suffix}"
                cursor.execute(
                    """
                    INSERT INTO cap_users
                      (employee_no, login_name, display_name, email, account_status, timezone_name)
                    VALUES (%s, %s, %s, %s, 'active', 'Asia/Shanghai')
                    """,
                    (f"EMP-HDR-{suffix}", f"user_{suffix}", entity_label, f"user_{suffix}@example.local"),
                )
                entity_type = "user"
                entity_id = int(cursor.lastrowid)
                result = {"ok": True, "user_id": entity_id, "affected_table": "cap_users"}

            elif screen_id == "roles-permissions":
                suffix = uuid.uuid4().hex[:6]
                entity_label = f"Frontend Role {suffix}"
                cursor.execute(
                    """
                    INSERT INTO cap_roles (role_code, role_name, description, data_scope, is_system_role, is_active)
                    VALUES (%s, %s, 'Created by screen primary action.', 'owned', 0, 1)
                    """,
                    (f"ROLE_HDR_{suffix.upper()}", entity_label),
                )
                entity_type = "role"
                entity_id = int(cursor.lastrowid)
                result = {"ok": True, "role_id": entity_id, "affected_table": "cap_roles"}

            elif screen_id == "field-config":
                suffix = uuid.uuid4().hex[:6]
                entity_label = f"Frontend Field {suffix}"
                cursor.execute(
                    """
                    INSERT INTO cap_custom_field_definitions
                      (entity_type, field_key, field_label, data_type, is_required, is_searchable, display_order, created_by)
                    VALUES ('project', %s, %s, 'text', 0, 1, 100, %s)
                    """,
                    (f"frontend_field_{suffix}", entity_label, user.user_id),
                )
                entity_type = "custom_field"
                entity_id = int(cursor.lastrowid)
                result = {"ok": True, "custom_field_id": entity_id, "affected_table": "cap_custom_field_definitions"}

            elif screen_id == "import-export":
                cursor.execute(
                    """
                    INSERT INTO cap_import_export_tasks
                      (task_code, task_kind, entity_type, task_status, source_file_uri, total_rows,
                       success_rows, failed_rows, requested_by, requested_at, tenant_id)
                    VALUES (%s, 'template_download', 'screen_primary', 'queued', NULL, 0, 0, 0, %s, %s, %s)
                    """,
                    (f"IET-HDR-{uuid.uuid4().hex[:8].upper()}", user.user_id, datetime.now(), tenant_of(user)),
                )
                entity_type = "import_export_task"
                entity_id = int(cursor.lastrowid)
                entity_label = "screen_primary"
                result = {"ok": True, "task_id": entity_id, "affected_table": "cap_import_export_tasks"}

            elif screen_id in {"research-library", "internal-research"}:
                suffix = uuid.uuid4().hex[:6]
                kind = "internal_report" if screen_id == "internal-research" else "market_clip"
                entity_label = f"Frontend Research {suffix}"
                cursor.execute(
                    """
                    INSERT INTO cap_research_notes
                      (note_code, note_kind, note_title, source_name, review_status, created_by, tenant_id)
                    VALUES (%s, %s, %s, 'Frontend', 'draft', %s, %s)
                    """,
                    (f"RES-HDR-{suffix.upper()}", kind, entity_label, user.user_id, tenant_of(user)),
                )
                entity_type = "research_note"
                entity_id = int(cursor.lastrowid)
                result = {"ok": True, "research_note_id": entity_id, "affected_table": "cap_research_notes"}

            elif screen_id == "manager-orgs":
                suffix = uuid.uuid4().hex[:6]
                entity_label = f"Frontend Manager {suffix}"
                cursor.execute(
                    """
                    INSERT INTO cap_management_orgs
                      (org_code, org_kind, org_name, city, status, tenant_id)
                    VALUES (%s, 'fund_manager', %s, 'Shanghai', 'active', %s)
                    """,
                    (f"MGR-HDR-{suffix.upper()}", entity_label, tenant_of(user)),
                )
                entity_type = "management_org"
                entity_id = int(cursor.lastrowid)
                result = {"ok": True, "management_org_id": entity_id, "affected_table": "cap_management_orgs"}

            elif screen_id == "burst-risk":
                suffix = uuid.uuid4().hex[:6]
                entity_label = f"Frontend Risk {suffix}"
                cursor.execute(
                    """
                    INSERT INTO cap_risk_incidents
                      (incident_code, incident_title, severity, incident_status, occurred_at, created_by, tenant_id)
                    VALUES (%s, %s, 'medium', 'open', %s, %s, %s)
                    """,
                    (f"RSK-HDR-{suffix.upper()}", entity_label, datetime.now(), user.user_id, tenant_of(user)),
                )
                entity_type = "risk_incident"
                entity_id = int(cursor.lastrowid)
                result = {"ok": True, "risk_incident_id": entity_id, "affected_table": "cap_risk_incidents"}

            elif screen_id == "post-data-collection":
                suffix = uuid.uuid4().hex[:6]
                entity_label = f"Frontend Campaign {suffix}"
                cursor.execute(
                    """
                    INSERT INTO cap_data_collection_campaigns
                      (campaign_code, campaign_name, period_code, status, due_on, frequency, send_mode, created_by, tenant_id)
                    VALUES (%s, %s, '2026Q3', 'draft', %s, 'quarterly', 'email', %s, %s)
                    """,
                    (f"DCC-HDR-{suffix.upper()}", entity_label, datetime.now().date(), user.user_id, tenant_of(user)),
                )
                entity_type = "data_collection_campaign"
                entity_id = int(cursor.lastrowid)
                result = {"ok": True, "collection_campaign_id": entity_id, "affected_table": "cap_data_collection_campaigns"}

            else:
                result = {"ok": True, "affected_table": "cap_ui_action_events", "screen_id": screen_id}

            audit_id = write_audit(
                cursor,
                user.user_id,
                action_code,
                entity_type,
                entity_id,
                entity_label,
                after=result,
            )
            result["audit_id"] = audit_id
            event_payload = GenericActionPayload(
                action=action_code,
                entity_type=entity_type,
                entity_id=entity_id,
                entity_label=entity_label,
                after=result,
            )
            result["event_id"] = write_ui_action_event(cursor, user.user_id, audit_id, event_payload, result)
        connection.commit()
    finally:
        connection.close()
    return result


@app.post("/api/actions")
def record_action(payload: GenericActionPayload, user: AuthedUser = Depends(current_user)) -> dict[str, Any]:
    connection = connect_db()
    try:
        with connection.cursor() as cursor:
            audit_id = write_audit(
                cursor,
                user.user_id,
                payload.action,
                payload.entity_type,
                payload.entity_id,
                payload.entity_label,
                before=payload.before,
                after=payload.after,
                risk_level=payload.risk_level,
            )
            result = {
                "ok": True,
                "audit_id": audit_id,
                "action": payload.action,
                "entity_type": payload.entity_type,
                "entity_id": payload.entity_id,
                "entity_label": payload.entity_label,
            }
            event_id = write_ui_action_event(cursor, user.user_id, audit_id, payload, result)
            result["event_id"] = event_id
        connection.commit()
    finally:
        connection.close()
    return result


@app.get("/api/ui-actions/recent")
def recent_ui_actions(user: AuthedUser = Depends(current_user)) -> dict[str, Any]:
    connection = connect_db()
    try:
        with connection.cursor() as cursor:
            ensure_ui_action_table(cursor)
            cursor.execute(
                """
                SELECT e.ui_action_event_id AS id, e.audit_log_id AS audit_id, u.display_name AS actor,
                       e.action_code, e.entity_type, e.entity_label, e.event_status, e.occurred_at,
                       e.payload_json, e.result_json
                FROM cap_ui_action_events e
                LEFT JOIN cap_users u ON u.user_id=e.actor_user_id
                WHERE e.tenant_id=%s
                ORDER BY e.ui_action_event_id DESC
                LIMIT 20
                """,
                (tenant_of(user),),
            )
            rows = cursor.fetchall()
        connection.commit()
    finally:
        connection.close()
    return {"count": len(rows), "items": rows}


@app.get("/api/audit/recent")
def recent_audit(user: AuthedUser = Depends(current_user)) -> dict[str, Any]:
    connection = connect_db()
    try:
        with connection.cursor() as cursor:
            cursor.execute(
                """
                SELECT a.audit_log_id AS id, u.display_name AS actor, a.action_code,
                       a.entity_type, a.entity_label, a.risk_level, a.occurred_at
                FROM cap_audit_logs a
                LEFT JOIN cap_users u ON u.user_id=a.actor_user_id
                WHERE a.tenant_id=%s
                ORDER BY a.audit_log_id DESC
                LIMIT 20
                """,
                (tenant_of(user),),
            )
            rows = cursor.fetchall()
    finally:
        connection.close()
    return {"count": len(rows), "items": rows}


# ── AI / LLM ──────────────────────────────────────────────────────────────
@app.get("/api/ai/status")
def ai_status(user: AuthedUser = Depends(current_user)) -> dict[str, Any]:
    """调试用:当前 LLM 是否配好、指向哪个 base_url/model(不回显 key)。"""
    return llm.status()


_BP_SYSTEM_PROMPT = (
    "你是 PE/VC 投资平台的 BP(商业计划书)解析助手。"
    "从用户提供的 BP 文本中抽取结构化字段,严格只返回一个 JSON 对象,不要额外解释。"
    "字段:short_name(企业简称)、legal_name(企业全称)、industry_group(行业方向)、"
    "city(城市)、funding_round(融资轮次)、summary(一句话摘要)、"
    "highlights(亮点数组)、risks(风险数组)、confidence(0-1 的整体置信度)。"
    "缺失字段用空字符串或空数组;不要编造事实。"
)


@app.post("/api/ai/parse-bp")
def parse_bp(payload: BpParsePayload, user: AuthedUser = Depends(current_user)) -> dict[str, Any]:
    """BP 文本 → 结构化项目字段(真模型)。可直接回填到「新增项目」表单。"""
    text = payload.text.strip()
    if not text:
        raise HTTPException(status_code=400, detail="text required")
    if len(text) > 20000:
        text = text[:20000]
    if not llm.is_configured():
        raise HTTPException(
            status_code=503,
            detail="AI 未配置:请在 deploy/.env 设置 LLM_ENABLED=true 及 LLM_BASE_URL / LLM_MODEL / LLM_API_KEY",
        )
    try:
        fields = llm.chat_json(
            [
                {"role": "system", "content": _BP_SYSTEM_PROMPT},
                {"role": "user", "content": f"BP 文本如下:\n\n{text}"},
            ],
            temperature=0.1,
        )
    except llm.LLMNotConfigured as exc:
        raise HTTPException(status_code=503, detail=str(exc))
    except llm.LLMError as exc:
        raise HTTPException(status_code=502, detail=str(exc))

    # 记一条真实的 AI 解析作业 + 审计(替代原来纯占位的 cap_ai_parse_jobs)。
    connection = connect_db()
    try:
        with connection.cursor() as cursor:
            cursor.execute(
                """
                INSERT INTO cap_ai_parse_jobs (job_code, job_name, parse_kind, requested_by, job_status)
                VALUES (%s, %s, 'business_plan', %s, 'completed')
                """,
                (f"AIP-{uuid.uuid4().hex[:8].upper()}", str(fields.get("short_name") or "BP 解析"), user.user_id),
            )
            job_id = int(cursor.lastrowid)
            write_audit(cursor, user.user_id, "ai.parse_bp", "ai_parse_job", job_id, str(fields.get("short_name") or ""), after={"model": llm.status()["model"]})
        connection.commit()
    finally:
        connection.close()

    return {"ok": True, "model": llm.status()["model"], "job_id": job_id, "fields": fields}


@app.post("/api/ai/summarize")
def ai_summarize(payload: AiSummarizePayload, user: AuthedUser = Depends(current_user)) -> dict[str, Any]:
    """会议纪要 / 研究材料 → 摘要 + 要点 + 待办(真模型)。"""
    text = payload.text.strip()
    if not text:
        raise HTTPException(status_code=400, detail="text required")
    if not llm.is_configured():
        raise HTTPException(status_code=503, detail="AI 未配置:请在 deploy/.env 配置 LLM_*")
    label = "会议纪要" if payload.kind == "meeting" else "研究材料"
    try:
        result = llm.chat_json(
            [
                {"role": "system", "content": f"你是投资平台助手。把{label}整理为 JSON:summary(摘要)、key_points(要点数组)、action_items(待办数组)。只返回 JSON。"},
                {"role": "user", "content": text[:20000]},
            ],
            temperature=0.2,
        )
    except llm.LLMNotConfigured as exc:
        raise HTTPException(status_code=503, detail=str(exc))
    except llm.LLMError as exc:
        raise HTTPException(status_code=502, detail=str(exc))
    return {"ok": True, "model": llm.status()["model"], "result": result}


# ── 真导入 / 导出(CSV)──────────────────────────────────────────────────
@app.get("/api/export/{screen_id}")
def export_ledger(screen_id: str, user: AuthedUser = Depends(current_user)) -> Response:
    """把某屏台账导成真实 CSV 文件(租户内,含 UTF-8 BOM 供 Excel 正确识别中文)。"""
    query = LEDGER_QUERIES.get(screen_id)
    if query is None:
        raise HTTPException(status_code=404, detail=f"No ledger mapping for {screen_id}")
    tenant_params: tuple[Any, ...] = (tenant_of(user),) if screen_id in LEDGER_TENANT_SCOPED else ()
    connection = connect_db()
    try:
        with connection.cursor() as cursor:
            cursor.execute(f"{query}\n            LIMIT 5000", tenant_params)
            rows = cursor.fetchall()
            write_audit(cursor, user.user_id, "export.csv", "screen", None, screen_id, after={"rows": len(rows)})
        connection.commit()
    finally:
        connection.close()

    buf = io.StringIO()
    if rows:
        writer = csv.DictWriter(buf, fieldnames=list(rows[0].keys()))
        writer.writeheader()
        for row in rows:
            writer.writerow({k: ("" if v is None else str(v)) for k, v in row.items()})
    body = ("﻿" + buf.getvalue()).encode("utf-8")  # BOM 让 Excel 认 UTF-8
    return Response(
        content=body,
        media_type="text/csv; charset=utf-8",
        headers={"Content-Disposition": f'attachment; filename="{screen_id}.csv"'},
    )


# 导入时的中英列名 → 项目字段映射(支持导出后原样回传)。
_IMPORT_PROJECT_COLS = {
    "short_name": "short_name", "项目名称": "short_name", "项目简称": "short_name",
    "legal_name": "legal_name", "企业全称": "legal_name",
    "industry_group": "industry_group", "行业方向": "industry_group",
    "city": "city", "城市": "city",
}


@app.post("/api/import/projects")
def import_projects(
    payload: ImportCsvPayload,
    user: AuthedUser = Depends(require_permission("project.edit")),
) -> dict[str, Any]:
    """CSV 批量建项目(真解析真入库,租户内 + project.edit)。逐行校验,返回成功数与错误行。"""
    reader = csv.DictReader(io.StringIO(payload.csv_text.lstrip("﻿")))
    created, errors = 0, []
    tid = tenant_of(user)
    connection = connect_db()
    try:
        with connection.cursor() as cursor:
            for i, raw in enumerate(reader, start=1):
                mapped: dict[str, str] = {}
                for col, val in raw.items():
                    key = _IMPORT_PROJECT_COLS.get((col or "").strip())
                    if key and val and val.strip():
                        mapped[key] = val.strip()
                if not mapped.get("short_name"):
                    errors.append({"row": i, "reason": "缺少项目名称/short_name"})
                    continue
                cursor.execute(
                    """
                    INSERT INTO cap_projects
                      (project_code, short_name, legal_name, opportunity_status, stage_label,
                       industry_group, city, owner_user_id, source_channel, summary, created_by, tenant_id)
                    VALUES (%s, %s, %s, 'sourced', 'Sourced', %s, %s, %s, 'CSV Import', 'Imported via CSV.', %s, %s)
                    """,
                    (
                        f"PRJ-CSV-{uuid.uuid4().hex[:8].upper()}",
                        mapped["short_name"],
                        mapped.get("legal_name") or mapped["short_name"],
                        mapped.get("industry_group", "Uncategorized"),
                        mapped.get("city", "-"),
                        user.user_id,
                        user.user_id,
                        tid,
                    ),
                )
                created += 1
            audit_id = write_audit(cursor, user.user_id, "import.projects", "project", None, f"CSV import {created} rows", after={"created": created, "errors": len(errors)})
        connection.commit()
    finally:
        connection.close()
    return {"ok": True, "created": created, "errors": errors, "audit_id": audit_id}


@app.post("/api/ai/analyze")
def ai_analyze(payload: AiAnalyzePayload, user: AuthedUser = Depends(current_user)) -> dict[str, Any]:
    """通用材料分析(AI 工作台):自定义指令 + 材料 → 自由文本分析(投资视角)。"""
    text = payload.text.strip()
    if not text:
        raise HTTPException(status_code=400, detail="text required")
    if not llm.is_configured():
        raise HTTPException(status_code=503, detail="AI 未配置:请在 deploy/.env 配置 LLM_*")
    instruction = (payload.instruction or "").strip() or "请对以下材料做投资视角的要点分析与风险提示。"
    try:
        answer = llm.chat(
            [
                {"role": "system", "content": "你是 PE/VC 投资平台的分析助手。用中文、结构化、可直接用于投资决策地回答;可用 Markdown 小标题与要点,避免空话。"},
                {"role": "user", "content": f"{instruction}\n\n材料:\n{text[:20000]}"},
            ],
            temperature=0.3,
            max_tokens=1500,
        )
    except llm.LLMNotConfigured as exc:
        raise HTTPException(status_code=503, detail=str(exc))
    except llm.LLMError as exc:
        raise HTTPException(status_code=502, detail=str(exc))
    return {"ok": True, "model": llm.status()["model"], "answer": answer}


@app.post("/api/ai/analyze/stream")
def ai_analyze_stream(payload: AiAnalyzePayload, user: AuthedUser = Depends(current_user)) -> StreamingResponse:
    """SSE 流式分析:边生成边下发 token,前端可低延迟增量渲染。"""
    text = payload.text.strip()
    if not text:
        raise HTTPException(status_code=400, detail="text required")
    if not llm.is_configured():
        raise HTTPException(status_code=503, detail="AI 未配置:请在 deploy/.env 配置 LLM_*")
    instruction = (payload.instruction or "").strip() or "请对以下材料做投资视角的要点分析与风险提示。"
    messages = [
        {"role": "system", "content": "你是 PE/VC 投资平台的分析助手。用中文、结构化、可直接用于投资决策地回答;可用 Markdown 小标题与要点,避免空话。"},
        {"role": "user", "content": f"{instruction}\n\n材料:\n{text[:20000]}"},
    ]

    def event_stream():
        try:
            for delta in llm.chat_stream(messages):
                yield f"data: {dumps({'delta': delta}, ensure_ascii=False)}\n\n"
            yield f"data: {dumps({'done': True})}\n\n"
        except (llm.LLMError, llm.LLMNotConfigured) as exc:
            yield f"data: {dumps({'error': str(exc)}, ensure_ascii=False)}\n\n"

    return StreamingResponse(
        event_stream(),
        media_type="text/event-stream",
        headers={"Cache-Control": "no-cache", "X-Accel-Buffering": "no", "Connection": "keep-alive"},
    )


# ── AI 配置(DB 覆盖 .env,管理端 UI 可改)──────────────────────────────────
_AI_SETTING_TO_ENV = {
    "llm_enabled": "LLM_ENABLED",
    "llm_base_url": "LLM_BASE_URL",
    "llm_api_key": "LLM_API_KEY",
    "llm_model": "LLM_MODEL",
    "llm_timeout": "LLM_TIMEOUT",
}


def _load_ai_settings_into_env() -> None:
    """把 DB 里的 LLM 配置写入 os.environ(覆盖 .env),供 llm.py 读取。启动 + 每次更新调用。"""
    if pymysql is None:
        return
    try:
        connection = connect_db()
    except Exception:  # noqa: BLE001 —— 启动时 DB 不可用不应崩
        return
    try:
        with connection.cursor() as cursor:
            cursor.execute("SELECT setting_key, setting_value FROM cap_app_settings WHERE setting_key LIKE 'llm_%%'")
            for row in cursor.fetchall():
                env = _AI_SETTING_TO_ENV.get(row["setting_key"])
                if env and row["setting_value"] is not None:
                    os.environ[env] = str(row["setting_value"])
    except Exception:  # noqa: BLE001 —— 表不存在等,忽略,退回 .env
        pass
    finally:
        connection.close()


@app.on_event("startup")
def _startup_load_ai_settings() -> None:
    _load_ai_settings_into_env()


@app.get("/api/ai/config")
def get_ai_config(user: AuthedUser = Depends(require_roles("system_admin", "managing_partner"))) -> dict[str, Any]:
    """管理端查看当前 AI 配置(绝不回显 key,只给 has_api_key)。"""
    return llm.status()


@app.post("/api/ai/config")
def set_ai_config(payload: AiConfigPayload, user: AuthedUser = Depends(require_roles("system_admin", "managing_partner"))) -> dict[str, Any]:
    """管理端更新 AI 配置 → 存 DB + 立即生效。api_key 留空表示不改。"""
    updates: dict[str, str] = {}
    if payload.enabled is not None:
        updates["llm_enabled"] = "true" if payload.enabled else "false"
    if payload.base_url is not None:
        updates["llm_base_url"] = payload.base_url.strip()
    if payload.api_key:
        updates["llm_api_key"] = payload.api_key.strip()
    if payload.model is not None:
        updates["llm_model"] = payload.model.strip()
    if payload.timeout is not None:
        updates["llm_timeout"] = str(payload.timeout)
    if not updates:
        raise HTTPException(status_code=400, detail="没有要更新的字段")
    connection = connect_db()
    try:
        with connection.cursor() as cursor:
            for key, val in updates.items():
                cursor.execute(
                    """
                    INSERT INTO cap_app_settings (setting_key, setting_value, updated_by)
                    VALUES (%s, %s, %s)
                    ON DUPLICATE KEY UPDATE setting_value=VALUES(setting_value), updated_by=VALUES(updated_by)
                    """,
                    (key, val, user.user_id),
                )
            write_audit(cursor, user.user_id, "ai.config.update", "app_setting", None, "LLM config", after={"keys": list(updates)})
        connection.commit()
    finally:
        connection.close()
    _load_ai_settings_into_env()  # 立即生效
    return {"ok": True, "config": llm.status()}


@app.post("/api/ai/config/test")
def test_ai_config(user: AuthedUser = Depends(require_roles("system_admin", "managing_partner"))) -> dict[str, Any]:
    """连通性测试:发一条极短提示,验证 base_url/model/key 是否可用。"""
    if not llm.is_configured():
        raise HTTPException(status_code=503, detail="AI 未启用或未配置")
    try:
        reply = llm.chat([{"role": "user", "content": "只回复两个字:正常"}], max_tokens=16)
    except llm.LLMError as exc:
        raise HTTPException(status_code=502, detail=str(exc))
    return {"ok": True, "model": llm.status()["model"], "reply": reply.strip()[:100]}


# ── 文件文本抽取(txt/md/pdf/docx → 纯文本,供 AI 分析)──────────────────────
def _extract_pdf(data: bytes) -> str:
    from pypdf import PdfReader

    reader = PdfReader(io.BytesIO(data))
    return "\n".join((page.extract_text() or "") for page in reader.pages[:200])


def _extract_docx(data: bytes) -> str:
    import docx

    document = docx.Document(io.BytesIO(data))
    parts = [p.text for p in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            parts.append(" | ".join(cell.text for cell in row.cells))
    return "\n".join(parts)


@app.post("/api/ai/extract-text")
async def extract_text(file: UploadFile = File(...), user: AuthedUser = Depends(current_user)) -> dict[str, Any]:
    """把上传文件抽成纯文本(txt/md 直读、pdf 用 pypdf、docx 用 python-docx),供 AI 分析。"""
    name = file.filename or "file"
    ext = os.path.splitext(name)[1].lower()
    data = await file.read(20 * 1024 * 1024 + 1)  # 20MB 上限
    await file.close()
    if len(data) > 20 * 1024 * 1024:
        raise HTTPException(status_code=413, detail="文件过大(上限 20MB)")
    try:
        if ext in (".txt", ".md", ".markdown"):
            text = data.decode("utf-8", "ignore")
        elif ext == ".pdf":
            text = _extract_pdf(data)
        elif ext == ".docx":
            text = _extract_docx(data)
        elif ext == ".doc":
            raise HTTPException(status_code=415, detail="旧版 .doc 不支持,请另存为 .docx 或 PDF")
        else:
            raise HTTPException(status_code=415, detail=f"不支持的文件类型:{ext or '未知'}")
    except HTTPException:
        raise
    except Exception as exc:  # noqa: BLE001 —— 抽取失败(损坏/加密等)统一提示
        raise HTTPException(status_code=422, detail=f"文件解析失败:{exc}")
    text = text.strip()[:50000]  # 抽取文本上限,防超模型上下文
    if not text:
        raise HTTPException(status_code=422, detail="未抽取到文本(可能是扫描件/图片型 PDF,需 OCR)")
    return {"ok": True, "file_name": name, "chars": len(text), "text": text}


# ── 真文件上传 / 下载 ──────────────────────────────────────────────────────
@app.post("/api/files/upload")
async def upload_file(file: UploadFile = File(...), user: AuthedUser = Depends(current_user)) -> dict[str, Any]:
    """真上传:落盘到 UPLOAD_DIR/tenant_<tid>/,返回可入库的 storage_uri。"""
    tid = tenant_of(user)
    safe = re.sub(r"[^\w.\-]", "_", os.path.basename(file.filename or "file"))[:120] or "file"
    rel = f"tenant_{tid}/{uuid.uuid4().hex}__{safe}"
    dest = os.path.join(UPLOAD_DIR, rel)
    os.makedirs(os.path.dirname(dest), exist_ok=True)
    size = 0
    try:
        with open(dest, "wb") as out:
            while True:
                chunk = await file.read(1024 * 1024)
                if not chunk:
                    break
                size += len(chunk)
                if size > MAX_UPLOAD_BYTES:
                    out.close()
                    os.remove(dest)
                    raise HTTPException(status_code=413, detail="文件过大(上限 50MB)")
                out.write(chunk)
    finally:
        await file.close()
    return {
        "ok": True,
        "storage_uri": f"file://{rel}",
        "file_name": file.filename or safe,
        "mime_type": file.content_type or "application/octet-stream",
        "size": size,
    }


@app.get("/api/files/download")
def download_file(uri: str, user: AuthedUser = Depends(current_user)) -> FileResponse:
    """下载:仅允许下本租户目录下的文件(防越权 / 路径穿越)。"""
    tid = tenant_of(user)
    if not uri.startswith("file://"):
        raise HTTPException(status_code=400, detail="bad uri")
    rel = uri[len("file://"):]
    if not rel.startswith(f"tenant_{tid}/") or ".." in rel:
        raise HTTPException(status_code=403, detail="无权访问该文件")
    path = os.path.join(UPLOAD_DIR, rel)
    if not os.path.isfile(path):
        raise HTTPException(status_code=404, detail="文件不存在")
    return FileResponse(path, filename=os.path.basename(rel).split("__", 1)[-1])


# ── 审批委托(cap_workflow_delegations)──────────────────────────────────────
def _same_tenant_user_ids_sql() -> str:
    # 同租户 = 组织树里 根公司(=tenant) 或其直属部门下的用户。
    return "org_id IN (SELECT org_id FROM cap_organizations WHERE org_id=%s OR parent_org_id=%s)"


@app.get("/api/users/simple")
def list_users_simple(user: AuthedUser = Depends(current_user)) -> dict[str, Any]:
    """同租户用户(供委托等选择)。"""
    tid = tenant_of(user)
    connection = connect_db()
    try:
        with connection.cursor() as cursor:
            cursor.execute(
                f"SELECT user_id AS id, display_name AS name, login_name FROM cap_users WHERE {_same_tenant_user_ids_sql()} AND account_status='active' AND deleted_at IS NULL ORDER BY display_name",
                (tid, tid),
            )
            rows = cursor.fetchall()
    finally:
        connection.close()
    return {"items": rows}


def _parse_dt(s: str | None, default: datetime) -> datetime:
    if not s:
        return default
    try:
        return datetime.fromisoformat(s.replace("Z", "").strip())
    except ValueError:
        return default


@app.get("/api/delegations")
def list_delegations(user: AuthedUser = Depends(current_user)) -> dict[str, Any]:
    """与我相关的委托(我委托出去的 + 委托给我的)。"""
    connection = connect_db()
    try:
        with connection.cursor() as cursor:
            cursor.execute(
                """
                SELECT d.delegation_id AS id, d.workflow_family, d.starts_at, d.ends_at, d.is_active, d.reason,
                       du.display_name AS delegator, eu.display_name AS delegatee,
                       d.delegator_user_id, d.delegatee_user_id
                FROM cap_workflow_delegations d
                LEFT JOIN cap_users du ON du.user_id=d.delegator_user_id
                LEFT JOIN cap_users eu ON eu.user_id=d.delegatee_user_id
                WHERE d.delegator_user_id=%s OR d.delegatee_user_id=%s
                ORDER BY d.delegation_id DESC
                """,
                (user.user_id, user.user_id),
            )
            rows = cursor.fetchall()
    finally:
        connection.close()
    return {"items": rows}


@app.post("/api/delegations")
def create_delegation(payload: DelegationPayload, user: AuthedUser = Depends(current_user)) -> dict[str, Any]:
    """设置审批委托:把某类工作流的审批权在一段时间内委托给同租户的另一人。"""
    if payload.delegatee_user_id == user.user_id:
        raise HTTPException(status_code=400, detail="不能委托给自己")
    fam = payload.workflow_family if payload.workflow_family in {"project", "fund", "office", "all"} else "all"
    now = datetime.now()
    starts = _parse_dt(payload.starts_at, now)
    ends = _parse_dt(payload.ends_at, now + timedelta(days=30))
    if ends <= starts:
        raise HTTPException(status_code=400, detail="结束时间必须晚于开始时间")
    tid = tenant_of(user)
    connection = connect_db()
    try:
        with connection.cursor() as cursor:
            # 校验受托人同租户
            cursor.execute(
                f"SELECT display_name FROM cap_users WHERE user_id=%s AND {_same_tenant_user_ids_sql()}",
                (payload.delegatee_user_id, tid, tid),
            )
            target = cursor.fetchone()
            if target is None:
                raise HTTPException(status_code=404, detail="受托人不存在或不在同一租户")
            cursor.execute(
                """
                INSERT INTO cap_workflow_delegations
                  (delegator_user_id, delegatee_user_id, workflow_family, starts_at, ends_at, is_active, reason)
                VALUES (%s, %s, %s, %s, %s, 1, %s)
                """,
                (user.user_id, payload.delegatee_user_id, fam, starts, ends, payload.reason),
            )
            did = int(cursor.lastrowid)
            audit_id = write_audit(cursor, user.user_id, "workflow.delegation.create", "delegation", did, target["display_name"], after={"family": fam})
        connection.commit()
    finally:
        connection.close()
    return {"ok": True, "delegation_id": did, "audit_id": audit_id}


@app.post("/api/delegations/{delegation_id}/revoke")
def revoke_delegation(delegation_id: int, user: AuthedUser = Depends(current_user)) -> dict[str, Any]:
    """撤销我发起的委托。"""
    connection = connect_db()
    try:
        with connection.cursor() as cursor:
            cursor.execute(
                "SELECT delegator_user_id FROM cap_workflow_delegations WHERE delegation_id=%s",
                (delegation_id,),
            )
            row = cursor.fetchone()
            if row is None:
                raise HTTPException(status_code=404, detail="委托不存在")
            if row["delegator_user_id"] != user.user_id:
                raise HTTPException(status_code=403, detail="只能撤销自己发起的委托")
            cursor.execute("UPDATE cap_workflow_delegations SET is_active=0 WHERE delegation_id=%s", (delegation_id,))
            audit_id = write_audit(cursor, user.user_id, "workflow.delegation.revoke", "delegation", delegation_id, str(delegation_id))
        connection.commit()
    finally:
        connection.close()
    return {"ok": True, "audit_id": audit_id}


@app.get("/api/db/ping")
def db_ping() -> dict[str, Any]:
    if pymysql is None:
        raise HTTPException(status_code=500, detail="pymysql is not installed")

    host = os.getenv("DB_HOST")
    user = os.getenv("DB_USER")
    password = os.getenv("DB_PASSWORD")
    database = os.getenv("DB_NAME", "capitalos")
    port = int(os.getenv("DB_PORT", "3306"))

    if not host or not user or password is None:
        raise HTTPException(status_code=400, detail="DB_HOST, DB_USER and DB_PASSWORD are required")

    connection = pymysql.connect(
        host=host,
        port=port,
        user=user,
        password=password,
        database=database,
        charset="utf8mb4",
        cursorclass=pymysql.cursors.DictCursor,
        connect_timeout=5,
    )
    try:
        with connection.cursor() as cursor:
            cursor.execute("SELECT COUNT(*) AS screen_count FROM cap_navigation_items")
            screen_count = cursor.fetchone()["screen_count"]
            cursor.execute(
                "SELECT COUNT(*) AS table_count FROM information_schema.tables "
                "WHERE table_schema=%s AND table_name LIKE 'cap_%%'",
                (database,),
            )
            table_count = cursor.fetchone()["table_count"]
    finally:
        connection.close()

    return {"ok": True, "database": database, "tables": table_count, "screens": screen_count}
